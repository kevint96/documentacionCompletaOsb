import streamlit as st
import os
import shutil
from zipfile import ZipFile
import zipfile
import tempfile
import subprocess
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.shared import Cm
import inspect
import os
import xml.etree.ElementTree as ET
import gspread
import time  # Importar el módulo time
import logging
import re
import inspect
import ast
from datetime import datetime
import difflib
import glob
import base64
import sys
import xml.etree.ElementTree as ET
from collections import defaultdict
from lxml import etree
import json
import zlib
import urllib.parse
import requests
import concurrent.futures

# URL del servidor público de PlantUML
PLANTUML_SERVER = "https://www.plantuml.com/plantuml/png/"
# Mapeo especial de caracteres para la codificación de PlantUML
PLANTUML_ENCODING = "0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz-_"

def print_with_line_number(msg):
    caller_frame = inspect.currentframe().f_back
    line_number = caller_frame.f_lineno
    st.success(f"Linea {line_number}: {msg}")
    print("")
    
def apply_format(run,fuente,size,negrita,color):
    run.font.name = fuente  # Cambiar el nombre de la fuente
    run.font.size = Pt(size)  # Cambiar el tamaño de la fuente
    run.font.bold = negrita  # Aplicar negrita
    run.font.color.rgb = RGBColor(0, 0, color)  # Cambiar el color del texto a rojo

def replace_text_in_paragraph(paragraph, replacements):
    full_text = paragraph.text
    contador = 1
    ##st.success(f"Texto en linea: {full_text}")
    for key, value in replacements.items():
        if key in full_text:
            ##st.success(f"full_text: {full_text}")
            ##st.success(f"p paragraphs: {paragraph.text}")
            ##st.success(f"clave coincide: {key}")
            full_text = full_text.replace(key, str(value))  # Actualiza full_text
            
            if key in '{nombre_servicio_inicial}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Arial',18,True,0)  # Aplicar formato al texto del párrafo
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                
            if key in '{nombre_operacion_inicial}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Arial',10,True,0)  # Aplicar formato al texto del párrafo
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                
            if key in '{nombre_servicio_secundario}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Arial',10,True,0)    # Aplicar formato al texto del párrafo
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            
            if key in '{nombre_operacion}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Times New Roman',10,False,0)    # Aplicar formato al texto del párrafo
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            if key in '{unique_operations}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Times New Roman',10,False,0)    # Aplicar formato al texto del párrafo
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            if key in '{nombre_servicio}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Times New Roman',10,False,0)    # Aplicar formato al texto del párrafo
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            if key in '{nombre_servicio_contrato}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Times New Roman',10,False,0)  # Aplicar formato al texto del párrafo
            
            if key in '{nombre_servicio_wsdl}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Times New Roman',10,False,0)  # Aplicar formato al texto del párrafo
            
            if key in '{nombre_servicio_contrato2}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Arial',10,False,0)  # Aplicar formato al texto del párrafo
                
            if key in '{nombre_servicio_tabla}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Arial MT',11,False,0)  # Aplicar formato al texto del párrafo
            
            if key in '{fecha}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Arial MT',10,False,0)  # Aplicar formato al texto del párrafo
            
            if key in '{autor_inicial}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Arial MT',10,True,0)  # Aplicar formato al texto del párrafo
            
            if key in '{autor}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Arial',10,False,0)  # Aplicar formato al texto del párrafo
            
            if key in '{autor2}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Arial MT',10,False,0)  # Aplicar formato al texto del párrafo
            
            if key in '{url}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Arial MT',10,False,255)  # Aplicar formato al texto del párrafo
                
            if key in '{operacion_legado}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Arial MT',10,False,255)  # Aplicar formato al texto del párrafo
                
            
            if key in '{proyecto_abc}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Arial MT',10,False,0)  # Aplicar formato al texto del párrafo

def print_element_content(element, element_name):
    #st.success(f"Contenido del {element_name}:")
    for paragraph in element.paragraphs:
        print(paragraph.text)
    for table in element.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    print(paragraph.text)

def replace_text_in_element(element, replacements):
    for paragraph in element.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)
    for table in element.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)

def replace_text_in_doc(doc, replacements):
    for p in doc.paragraphs:
        replace_text_in_paragraph(p, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)

    for section in doc.sections:
        #st.success(f"Encabezado de la sección: {section.header}")
        #print_element_content(section.header, "Encabezado de la sección")
        replace_text_in_element(section.header, replacements)
        #st.success(f"Pie de página de la sección: {section.footer}")
        #print_element_content(section.footer, "Pie de página de la sección")
        replace_text_in_element(section.footer, replacements)
        # Agregamos este bloque específico para procesar las tablas dentro del encabezado de la sección 2
        if "Encabezado-Sección 2-" in [paragraph.text for paragraph in section.header.paragraphs]:
            # for table in section.header.tables:
                # for row in table.rows:
                    # for cell in row.cells:
                        # for paragraph in cell.paragraphs:
                            # print(paragraph.text)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, replacements)
    
    doc = reemplazar_texto_en_doc(doc, replacements)
    
    return doc

def reemplazar_texto_con_imagen(doc_path, img_path, marcador="{Imagen_diagrama}"):
    doc = Document(doc_path)

    for para in doc.paragraphs:
        if marcador in para.text:
            # Borrar el texto del marcador
            para.text = para.text.replace(marcador, "")

            # Insertar la imagen en el mismo párrafo
            run = para.add_run()
            run.add_picture(img_path, width=Inches(6))  # Ajusta el tamaño si es necesario
            break  # Detenernos después de la primera coincidencia

    doc.save(doc_path)
    print(f"Diagrama insertado en: {doc_path}")
   
def service_refs_ruta_pipeline(pipeline_path, project_path):
    
    elemento = ""
    # Servicios a excluir
    servicios_a_excluir = [
        'ComponentesComunes/Proxies/PS_ManejadorGenericoErroresV1.0',
        'UtilitariosEBS/Proxies/AuditoriaSOA/RegistrarAuditoriaSOADATV1.0'
    ]
    
    while True:
        
        #st.success(f"pipeline_path: {pipeline_path}")
        
        # Leer el archivo .pipeline
        with open(pipeline_path, 'r') as file:
            pipeline_content = file.readlines()

        # Buscar todas las líneas que contienen ':service ref="'
        matching_lines = [line for line in pipeline_content if ':service ref="' in line]

        # Extraer la información deseada de las líneas coincidentes
        servicios = set()  # Usamos un conjunto para evitar elementos duplicados
        for line in matching_lines:
            service_start_index = line.find(':service ref="') + len(':service ref="')
            service_end_index = line.find('"', service_start_index)
            service_ref = line[service_start_index:service_end_index]
            # Verificar si el servicio no está en la lista de servicios a excluir
            if service_ref not in servicios_a_excluir:
                servicios.add(service_ref)

        # Imprimir los servicios encontrados
        st.success("Servicios encontrados:")
        for service in servicios:
            st.success(service)
            
             # Si el elemento contiene '/BusinessServices/', salir del bucle
            if '/BusinessServices/' in service:
                #st.success(f"BusinessServices: {service}")
                business_path = os.path.join(project_path, service + '.bix')
                
                with open(business_path, 'r') as business_file:
                    business_content = business_file.readlines()
                    
                    matching_lines = [line for line in business_content if 'operation-name>' in line]
                    
                    # Extraer los elementos ref de las líneas coincidentes
                    elementos_ref = set()  # Usamos un conjunto para evitar elementos duplicados
                    for line in matching_lines:
                        invoke_start_index = line.find('operation-name>') + len('operation-name>')
                        invoke_end_index = line.find('<', invoke_start_index)
                        invoke_ref = line[invoke_start_index:invoke_end_index]
                        elementos_ref.add(invoke_ref)

                    # Imprimir los elementos ref encontrados
                    st.success("Elementos ref encontrados en {}: ".format(service))
                    for elemento in elementos_ref:
                        st.success(elemento)
                return elemento

            # Construir la ruta del archivo proxy
            proxy_path = os.path.join(project_path, service + '.proxy')

            # Verificar si el archivo proxy existe
            if os.path.exists(proxy_path):
                # Leer el archivo proxy
                with open(proxy_path, 'r') as proxy_file:
                    proxy_content = proxy_file.readlines()

                # Buscar todas las líneas que contienen ':invoke ref="'
                matching_lines = [line for line in proxy_content if ':invoke ref="' in line]

                # Extraer los elementos ref de las líneas coincidentes
                elementos_ref = set()  # Usamos un conjunto para evitar elementos duplicados
                for line in matching_lines:
                    invoke_start_index = line.find(':invoke ref="') + len(':invoke ref="')
                    invoke_end_index = line.find('"', invoke_start_index)
                    invoke_ref = line[invoke_start_index:invoke_end_index]
                    elementos_ref.add(invoke_ref)

                # Imprimir los elementos ref encontrados
                st.success("Elementos ref encontrados en {}: ".format(service))
                for elemento in elementos_ref:
                    st.success(elemento)

                    # Si el elemento contiene '/BusinessServices/', salir del bucle
                    if '/BusinessServices/' in elemento:
                        #st.success(f"elemento: {elemento}")
                        return elemento
                    else:
                        pipeline_path = os.path.join(project_path, elemento + '.pipeline')
                       
            else:
                st.success("El archivo proxy {} no existe.".format(proxy_path))
                break

    return elemento

def extract_xsd_import_paths(wsdl_path):
    xsd_import_paths = set()  # Conjunto para almacenar rutas únicas

    # Leer el contenido del archivo WSDL
    with open(wsdl_path, 'r', encoding='utf-8') as file:
        wsdl_content = file.read()

    # Extraer el contenido dentro de CDATA usando una expresión regular
    cdata_match = re.search(r'<!\[CDATA\[(.*?)\]\]>', wsdl_content, re.DOTALL)
    
    if cdata_match:
        cdata_content = cdata_match.group(1)  # Obtener solo el contenido dentro de CDATA

        # Parsear el contenido XML dentro del CDATA
        try:
            root = ET.fromstring(cdata_content)  # Convertir el CDATA en XML
        except ET.ParseError as e:
            print("Error al parsear el contenido de CDATA:", e)
            return xsd_import_paths

        # Espacios de nombres comunes en WSDL
        namespaces = {
            'xsd': 'http://www.w3.org/2001/XMLSchema'
        }

        # Buscar todos los elementos <xsd:import>
        for xsd_import in root.findall(".//xsd:import", namespaces):
            schema_location = xsd_import.get("schemaLocation")
            if schema_location:
                xsd_import_paths.add(schema_location)
    return list(xsd_import_paths)  # Convertimos el conjunto de vuelta a lista antes de devolverlo

def find_import_elements_with_namespace(xsd_content, target_namespace, xsd_file_path):
    schema_location = ""
    absolute_schema_location = None  # Inicializa la variable

    namespaces = {
        'xsd': 'http://www.w3.org/2001/XMLSchema'
        # Agrega otros namespaces si es necesario
    }
    #st.success(f"target_namespace: {target_namespace}")

    root = ET.fromstring(xsd_content)
    
    #st.success(f"xsd_file_path: {xsd_file_path}")
    
    # Busca todos los elementos import
    xsd_import_elements = root.findall(".//xsd:import", namespaces)

    for import_element in xsd_import_elements:
        namespace = import_element.get('namespace')
        #st.success(f"namespace: {namespace}")
        if namespace == target_namespace:
            schema_location = import_element.get('schemaLocation')
            #st.success(f"Found xsd:import with namespace '{namespace}': {schema_location}")
            
            # Concatena la ruta del archivo XSD principal con la ubicación del esquema importado
            absolute_schema_location = os.path.normpath(os.path.join(os.path.dirname(xsd_file_path), schema_location)).replace('\\', '/')
            #st.success(f"schema_location: {absolute_schema_location}")
            break  # Si encuentras la coincidencia, sal del bucle
    
    return absolute_schema_location  # Esto devolverá None si no se encontró coincidencia "

def extract_namespaces(xsd_content):
    """Extrae los namespaces definidos en el XSD."""
    namespaces = {}
    matches = re.findall(r'xmlns:([\w]+)="([^"]+)"', xsd_content)
    for prefix, uri in matches:
        namespaces[prefix] = uri
    return namespaces


def extract_imports(root):
    """Extrae los imports y los mapea con sus schemaLocation."""
    # Detectar el prefijo correcto para XML Schema (puede ser xs: o xsd:)
    schema_ns = "http://www.w3.org/2001/XMLSchema"
    prefix = None
    
    # Buscar en los atributos del root el namespace correspondiente
    for attr in root.attrib:
        if root.attrib[attr] == schema_ns:
            prefix = attr.split(":")[-1]  # Extraer el prefijo después de "xmlns:"
            break
    
    # Si no se encontró prefijo, usar xs por defecto
    if not prefix:
        prefix = "xs"

    # Buscar los imports con el prefijo detectado
    imports = {}
    for imp in root.findall(f".//{prefix}:import", {prefix: schema_ns}):
        namespace = imp.attrib.get('namespace')
        schema_location = imp.attrib.get('schemaLocation')
        if namespace and schema_location:
            imports[namespace] = schema_location
    
    return imports

def get_correct_xsd_path(current_xsd_path, schema_location):
    """
    Corrige la ruta de un XSD importado considerando los niveles de directorio.
    """
    base_path = os.path.dirname(current_xsd_path)  # Obtener la carpeta del XSD actual
    #st.success(f"base_path: {base_path}")
    corrected_path = os.path.abspath(os.path.join(base_path, schema_location))
    #st.success(f"corrected_path: {corrected_path}")    # Resolver la ruta correcta
    corrected_path = corrected_path.replace("/mount/src/documentacioncompletaosb/extraccion_jar","")
    corrected_path = corrected_path.replace("/mount/src/documentacioncompletaosb","")
    corrected_path = corrected_path.replace(".xsd",".XMLSchema")
    #st.success(f"corrected_path: {corrected_path}")    # Resolver la ruta correcta

    return corrected_path

def parse_xsd_file(project_path, xsd_file_path, operation_name, service_url, capa_proyecto, 
                   operacion_business, operations, service_name, operation_actual, 
                   target_complex_type=None, root_element_name=None,
                   request_elements=None, response_elements=None,processed_types=None):
    """
    Parsea un XSD y extrae los elementos request/response de forma recursiva.
    """

    # 🔹 Asegurar que las listas no se reinicien
    if request_elements is None:
        request_elements = []
    if response_elements is None:
        response_elements = []
    if processed_types is None:
        processed_types = set()  # 🔹 Inicializa el conjunto solo si no existe
        processed_types = {}

    extraccion_dir = os.path.abspath(project_path)
    xsd_file_path = os.path.normpath(xsd_file_path.strip("/\\"))  
    subcarpeta_xsd = os.path.dirname(xsd_file_path)
    subcarpeta_xsd = os.path.normpath(subcarpeta_xsd).replace("../", "")

    ruta_corregida = os.path.join(extraccion_dir, subcarpeta_xsd, os.path.basename(xsd_file_path))
    
    #print_with_line_number(f"extraccion_dir: {extraccion_dir}")
    #print_with_line_number(f"xsd_file_path: {xsd_file_path}")
    #print_with_line_number(f"subcarpeta_xsd: {subcarpeta_xsd}")
    #print_with_line_number(f"Ruta corregida FINAL: {ruta_corregida}")
    
    if not os.path.isfile(ruta_corregida):
        st.error(f"El archivo XSD {ruta_corregida} no existe.")
        return request_elements, response_elements

    # Leer el contenido del XSD
    try:
        with open(ruta_corregida, 'r', encoding="utf-8") as f:
            xsd_content = f.read()
    except FileNotFoundError:
        return None
    # Extraer el contenido de CDATA si es necesario
    cdata_match = re.search(r'<!\[CDATA\[(.*?)\]\]>', xsd_content, re.DOTALL)
    if cdata_match:
        xsd_content = cdata_match.group(1)
        #print_with_line_number("Se ha extraído el contenido de CDATA correctamente")

    try:
        root = ET.fromstring(xsd_content)
    except ET.ParseError as e:
        st.error(f"Error al analizar el XMLSchema: {e}")
        return request_elements, response_elements

    namespaces = extract_namespaces(xsd_content)
    imports = extract_imports(root)

    #print_with_line_number(f"Namespaces detectados: {namespaces}")
    #print_with_line_number(f"Imports encontrados: {imports}")
    
    # for item in request_elements:
        # print_with_line_number(f"item['name'] request: {item['name']}")
    
    # for item in response_elements:
        # print_with_line_number(f"item['name'] response: {item['name']}")

    # 🔹 Verificar qué prefijos están en el namespaces
    valid_prefixes = [p for p in ['xs', 'xsd'] if p in namespaces]

    if not valid_prefixes:
        st.error("⛔ No se encontró un prefijo válido en los namespaces del XSD")
        return request_elements, response_elements  # Salir si no hay prefijos válidos

    # 🔹 Tomar el primer prefijo encontrado en namespaces (xs o xsd)
    prefix = valid_prefixes[0]
    #print_with_line_number(f"prefix: {prefix}")

    # 🔹 Buscar complexTypes con el prefijo detectado dinámicamente
    complex_types = {
        elem.attrib.get('name', None): elem
        for elem in root.findall(f".//{prefix}:complexType", namespaces)
        if 'name' in elem.attrib
    }

    # 🔹 Buscar todos los elementos principales con el prefijo detectado
    root_elements = {
        elem.attrib.get('name', ''): elem.attrib.get('type', '').split(':')[-1]
        for elem in root.findall(f".//{prefix}:element", namespaces)
    }

    # 🚀 **Si `target_complex_type` está definido, buscar SOLO ese complexType.**
    if target_complex_type:
        #print_with_line_number(f"🔍 Buscando SOLO el complexType: {target_complex_type}")
        explorar_complex_type(target_complex_type, root_element_name, complex_types, namespaces, imports, extraccion_dir, 
                              xsd_file_path, project_path, service_url, capa_proyecto, operacion_business, 
                              operations, service_name, operation_actual, request_elements, response_elements, operation_name,processed_types)
        return request_elements, response_elements

    # 🔹 Si `target_complex_type` no está, procesamos TODO desde los elementos raíz.
    for root_element_name, complex_type in root_elements.items():
        #print_with_line_number(f"Procesando raíz: {root_element_name} -> {complex_type}")

        if complex_type in complex_types:
            explorar_complex_type(complex_type, root_element_name, complex_types, namespaces, imports, extraccion_dir, 
                                  xsd_file_path, project_path, service_url, capa_proyecto, operacion_business, 
                                  operations, service_name, operation_actual, request_elements, response_elements, operation_name,processed_types)

    print_with_line_number(f"Total elementos request: {len(request_elements)}")
    print_with_line_number(f"Total elementos response: {len(response_elements)}")
    return request_elements, response_elements


def explorar_complex_type(type_name, parent_element_name, complex_types, namespaces, imports, extraccion_dir, 
                          xsd_file_path, project_path, service_url, capa_proyecto, operacion_business, 
                          operations, service_name, operation_actual, request_elements, response_elements, operation_name,processed_types=None):
    """Explora recursivamente un complexType y extrae sus elementos internos."""

    if processed_types is None:
        #processed_types = set()
        processed_types = {}

    type_name = type_name.split(':')[-1]  

    # if type_name in processed_types:
        # print_with_line_number(f"🔄 parent_element_name: {parent_element_name}")
        # print_with_line_number(f"🔄 Se detectó recursión en {type_name}, evitando ciclo infinito.")
        
        # element_details = {
                    # 'elemento': parent_element_name.split('.')[0],  
                    # 'name': parent_element_name,  
                    # 'type': type_name,
                    # 'url': service_url,
                    # 'ruta': capa_proyecto,
                    # 'minOccurs': 0,
                    # 'operations': operations,
                    # 'service_name': service_name,
                    # 'operation_actual': operation_actual,
                # }
        # print_with_line_number(f"element_details: {element_details}")

        # if 'Request' in parent_element_name:
            # request_elements.append(element_details)
        # elif 'Response' in parent_element_name:
            # response_elements.append(element_details)
        # return  # Evita seguir procesando un tipo ya visitado
    
    processed_types.setdefault(type_name, [])  # ✅ Registrar que ya se visitó este tipo

    if type_name in complex_types:
        #print_with_line_number(f"Explorando complexType: {type_name}")

        # 🔹 Buscar un prefijo válido
        prefix = next((p for p in ['xs', 'xsd'] if p in namespaces), None)
        if not prefix:
            st.error(f"⛔ No se encontró un prefijo válido en namespaces: {namespaces}")
            return
        
        # 🔹 Buscar 'sequence' con prefijo válido
        sequence = complex_types[type_name].find(f'{prefix}:sequence', namespaces)
        if sequence is None:
            #st.warning(f"⚠ No se encontró 'sequence' en {type_name}")
            
            complex_content = complex_types[type_name].find(f'{prefix}:complexContent', namespaces)
            if complex_content is not None:
                extension = complex_content.find(f'{prefix}:extension', namespaces)
                if extension is not None and 'base' in extension.attrib:
                    base_type = extension.attrib['base'].split(":")[-1]  # Obtener el nombre sin prefijo
                    
                    #print_with_line_number(f"🔄 {type_name} extiende {base_type}, explorando {base_type}...")
                    explorar_complex_type(base_type, parent_element_name, complex_types, namespaces, imports, 
                                          extraccion_dir, xsd_file_path, project_path, service_url, capa_proyecto, 
                                          operacion_business, operations, service_name, operation_actual, 
                                          request_elements, response_elements, operation_name,processed_types)
                    return  # Salimos porque ya delegamos la exploración a la base
                
            #st.warning(f"⚠ No se encontró ni 'sequence' ni 'extension' en {type_name}")
            return  # Si no hay ni sequence ni extensión, no hay nada más que hacer

        ##print_with_line_number(f"Usando prefijo: {prefix}")

        if prefix not in namespaces:
            st.error(f"⛔ Error: el prefijo '{prefix}' no está en namespaces: {namespaces}")
            return

        for element in sequence.findall(f'{prefix}:element', namespaces):
            element_name = element.attrib.get('name', '')
            element_type = element.attrib.get('type', '')
            element_minOccurs = element.attrib.get('minOccurs', '')
            if element_minOccurs is None:
                element_minOccurs = 0
           
            #print_with_line_number(f"element_name: {element_name}")
            #print_with_line_number(f"element_type: {element_type}")
            #print_with_line_number(f"element_minOccurs: {element_minOccurs}")
            full_name = f"{parent_element_name}.{element_name}" if parent_element_name else element_name
            print_with_line_number(f"Encontrado elemento: {full_name}")
            #print_with_line_number(f"Encontrado elemento: {full_name} con tipo: {element_type} y minOcurs: {element_minOccurs}")

            # 🔹 Buscar 'simpleType' con prefijo válido
            simple_type = element.find(f'{prefix}:simpleType', namespaces)
            if simple_type is not None:
                restriction = simple_type.find(f'{prefix}:restriction', namespaces)
                if restriction is not None and 'base' in restriction.attrib:
                    element_type = restriction.attrib['base']
                    #print_with_line_number(f"Elemento {full_name} tiene restricción con base: {element_type}")

            if element_type.startswith(("xsd:", "xs:")):
                element_details = {
                    'elemento': parent_element_name.split('.')[0],  
                    'name': full_name,  
                    'type': element_type,
                    'url': service_url,
                    'ruta': capa_proyecto,
                    'minOccurs': element_minOccurs,
                    'operations': operations,
                    'service_name': service_name,
                    'operation_actual': operation_actual,
                }
                #print_with_line_number(f"Agregando elemento primitivo: {element_details}")

                if 'Request' in parent_element_name:
                    request_elements.append(element_details)
                elif 'Response' in parent_element_name:
                    response_elements.append(element_details)

            elif element_type in complex_types:
                #print_with_line_number(f"Buscando {element_type} en el mismo XSD")
                explorar_complex_type(element_type, full_name, complex_types, namespaces, imports, extraccion_dir, 
                                      xsd_file_path, project_path, service_url, capa_proyecto, operacion_business, 
                                      operations, service_name, operation_actual, request_elements, response_elements, operation_name,processed_types)

            elif ':' in element_type:
                prefix, nested_type = element_type.split(':')
                
                if nested_type in complex_types:
                    #print_with_line_number(f"Buscando {nested_type} en el mismo XSD")
                    explorar_complex_type(nested_type, full_name, complex_types, namespaces, imports, extraccion_dir, 
                                          xsd_file_path, project_path, service_url, capa_proyecto, operacion_business, 
                                          operations, service_name, operation_actual, request_elements, response_elements, operation_name,processed_types)
                elif prefix in namespaces:
                    namespace = namespaces[prefix]
                    if namespace in imports:
                        schema_location = imports[namespace]
                        #st.warning(f"El tipo {nested_type} está en otro XSD: {schema_location}")
                        corrected_xsd_path = get_correct_xsd_path(xsd_file_path, schema_location)
                        #print_with_line_number(f"corrected_xsd_path: {corrected_xsd_path}")
                        new_xsd_path = os.path.join(extraccion_dir, corrected_xsd_path)
                        #print_with_line_number(f"new_xsd_path: {new_xsd_path}")

                        parse_xsd_file(project_path, new_xsd_path, operation_name, service_url, 
                                       capa_proyecto, operacion_business, operations, 
                                       service_name, operation_actual, 
                                       target_complex_type=nested_type, 
                                       root_element_name=full_name,
                                       request_elements=request_elements,
                                       response_elements=response_elements,
                                       processed_types=processed_types)
                    else:
                        st.warning(f"No se encontró el namespace para el prefijo {prefix}")
                else:
                    st.warning(f"complexType {element_type} no encontrado en el XSD")
    else:
        st.warning(f"complexType {type_name} no encontrado en el XSD")

def leer_xsd_file(xsd_file_path, complexType_name):
    elements_list = []

    if xsd_file_path.endswith('.xsd') and os.path.isfile(xsd_file_path):
        with open(xsd_file_path, 'r', encoding="utf-8") as f:
            xsd_content = f.read()
            root = ET.fromstring(xsd_content)
            namespaces = {'xs': 'http://www.w3.org/2001/XMLSchema'}
            
            #st.success(f"xsd_file_path: {xsd_file_path}")
            

            # Función para detectar y eliminar repeticiones cíclicas en los nombres de los elementos
            def remove_repetitions(element_name):
                parts = element_name.split('.')
                seen = set()
                unique_parts = []
                for part in parts:
                    if part in seen:
                        break
                    seen.add(part)
                    unique_parts.append(part)
                return '.'.join(unique_parts)

            # Función para obtener elementos recursivamente con control de visitas
            def get_elements(complex_type_element, parent_name, visited):
                sequence_element = complex_type_element.find('xs:sequence', namespaces)
                if sequence_element is not None:
                    child_elements = sequence_element.findall('xs:element', namespaces)
                    for child_element in child_elements:
                        element_name = child_element.attrib.get('name', '')
                        element_type = child_element.attrib.get('type', '')
                        full_element_name = f"{parent_name}.{element_name}"

                        # Detectar y eliminar repeticiones cíclicas
                        full_element_name = remove_repetitions(full_element_name)

                        #st.success(f"element_name: {full_element_name}")
                        #st.success(f"element_type: {element_type}")
                        if not element_type:
                            element_type = 'xs:string'
                        elements_list.append({'element_name': full_element_name, 'element_type': element_type})

                        if ':' in element_type:
                            prefix, complexType_name_interno = element_type.split(':')
                            if complexType_name_interno not in visited:
                                visited.add(complexType_name_interno)
                                complex_type_element = root.find(f".//xs:complexType[@name='{complexType_name_interno}']", namespaces)
                                if complex_type_element is not None:
                                    get_elements(complex_type_element, full_element_name, visited)

            complex_type_element = root.find(f".//xs:complexType[@name='{complexType_name}']", namespaces)
            if complex_type_element is not None:
                
                #st.success(f"complex_type_name: {complexType_name}")
                
                #st.success(f"complex_type_element: {complex_type_element}")
                
                
                visited = set()
                get_elements(complex_type_element, complexType_name, visited)
                
    return elements_list
    
def has_http_provider_id(xml_content):
    root = ET.fromstring(xml_content)
    namespaces = {'tran': 'http://www.bea.com/wli/sb/transports'}
    provider_id_element = root.find(".//tran:provider-id", namespaces)
    return provider_id_element is not None and provider_id_element.text == 'http'

def extract_project_name_from_proxy(proxy_path):
    try:
        with open(proxy_path, 'r', encoding="utf-8") as f:
            content = f.read()
            start = content.find('<con:wsdl ref="') + len('<con:wsdl ref="')
            end = content.find('"', start)
            wsdl_ref = content[start:end]
            return wsdl_ref.split("/")[0]
    except FileNotFoundError:
        ##st.success(f"El archivo {proxy_path} no existe.")
        return None

def reemplazar_texto_en_doc(doc, reemplazos):
    """
    Reemplaza variables en el documento, incluyendo encabezados, pies de página y contenido.
    """
    # Reemplazo en párrafos normales
    for parrafo in doc.paragraphs:
        for clave, valor in reemplazos.items():
            if clave in parrafo.text:
                parrafo.text = parrafo.text.replace(clave, valor)
    
    # Reemplazo en encabezados y pies de página
    for section in doc.sections:
        # Encabezado
        for parrafo in section.header.paragraphs:
            for clave, valor in reemplazos.items():
                if clave in parrafo.text:
                    parrafo.text = parrafo.text.replace(clave, valor)
        
        # Pie de página
        for parrafo in section.footer.paragraphs:
            for clave, valor in reemplazos.items():
                if clave in parrafo.text:
                    parrafo.text = parrafo.text.replace(clave, valor)
    
    # Reemplazo en tablas sin alterar el formato
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for clave, valor in reemplazos.items():
                    if clave in celda.text:
                        celda.text = celda.text.replace(clave, valor)
    
    return doc

def extract_service_url(xml_content):
    root = ET.fromstring(xml_content)
    tran_namespace = {'tran': 'http://www.bea.com/wli/sb/transports', 'env': 'http://www.bea.com/wli/config/env'}
    uri_element = root.find(".//tran:URI/env:value", namespaces=tran_namespace)
    if uri_element is not None:
        return uri_element.text
    return ''

def extract_pipeline_path_from_proxy(proxy_path, jdeveloper_projects_dir):
    try:
        with open(proxy_path, 'r', encoding="utf-8") as f:
            content = f.read()
            start = content.find('<ser:invoke ref="') + len('<ser:invoke ref="')
            end = content.find('"', start)
            pipeline_ref = content[start:end]
            pipeline_path = os.path.join(jdeveloper_projects_dir, pipeline_ref + ".Pipeline")
            return pipeline_path
    except FileNotFoundError:
        print(f"El archivo {proxy_path} no pudo ser encontrado.")
        return None  # O puedes lanzar otra excepción, dependiendo del flujo de tu programa.
     
def extract_wsdl_relative_path(xml_content):
    root = ET.fromstring(xml_content)
    namespaces = {'con': 'http://www.bea.com/wli/sb/services/bindings/config'}
    wsdl_ref_element = root.find(".//con:wsdl", namespaces)
    if wsdl_ref_element is not None:
        wsdl_relative_path = wsdl_ref_element.attrib.get('ref', '')
        return wsdl_relative_path
    return ''
    
def extract_wsdl_operations(wsdl_path):
    operations = set()  # Utilizamos un conjunto en lugar de una lista
    if wsdl_path.endswith('.WSDL') and os.path.isfile(wsdl_path):
        with open(wsdl_path, 'r', encoding="utf-8") as f:
            wsdl_content = f.read()
            # Buscamos todas las coincidencias de "<operation name=" seguidas por el nombre de la operación
            operation_names = re.findall(r'operation name="([^"]+)', wsdl_content)
            for operation_name in operation_names:
                operations.add(operation_name)  # Agregamos el nombre de la operación al conjunto
    return list(operations)  # Convertimos el conjunto de vuelta a lista antes de devolverlo

def extraer_operaciones_expuestas_http(project_path):
    wsdl_operations_map = {}
    for root, dirs, files in os.walk(project_path):
        if os.path.basename(root) == "Proxies":
            ##st.success(f"✅ Proxies {elementos_xsd}")
            for file in files:
                if file.endswith('.ProxyService'):
                    osb_file_path = os.path.join(root, file)
                    #st.success(f"✅ file {file}")
                    #st.success(f"✅ osb_file_path {osb_file_path}")
                    project_name = extract_project_name_from_proxy(osb_file_path)
                    
                    if project_name is None:
                        continue 
                    pipeline_path = extract_pipeline_path_from_proxy(osb_file_path, project_path)
                    ##st.success(f"✅ pipeline_path {pipeline_path}")
                    with open(osb_file_path, 'r', encoding="utf-8") as f:
                        content = f.read()
                        if "EBS" not in project_name and "ABC" not in project_name:
                            if has_http_provider_id(content):
                                service_name = os.path.splitext(file)[0]
                                #st.success(f"✅ project_name {project_name}")
                                #st.success(f"✅ service_name {service_name}")
                                service_url = extract_service_url(content)
                                #st.success(f"✅ service_url {service_url}")
                                wsdl_relative_path = extract_wsdl_relative_path(content)
                                if wsdl_relative_path:
                                    wsdl_path = os.path.join(project_path, wsdl_relative_path + ".WSDL")
                                    capa_proyecto = '/'+ wsdl_relative_path.split('/')[0]
                                    
                                    #st.success(f"capa_proyecto: {capa_proyecto}")
                                    
                                    #st.success(f"wsdl_path: {wsdl_path}")
                                    operations = extract_wsdl_operations(wsdl_path)
                                    wsdl_operations_map[wsdl_path] = (
                                        operations, project_name, service_name, osb_file_path,pipeline_path, service_url, capa_proyecto
                                    )
    #st.success(f"✅ wsdl_operations_map {wsdl_operations_map}")
    return wsdl_operations_map

def extraer_schemas_operaciones_expuestas_http(project_path,operacion_a_documentar):
    
    osb_services = []
    elementos_xsd = []
    operations =[]
    operation_to_xsd = {}
    services_for_operations = {}
    found = False  # Variable para rastrear si se encuentra la operación

    #print_with_line_number(f"URL generada: {url}")
    wsdl_operations_map = extraer_operaciones_expuestas_http(project_path)
    
    # Recorriendo el diccionario
    for wsdl_path, data in wsdl_operations_map.items():
        # Desempaquetar la tupla
        operations, project_name, service_name, osb_file_path, pipeline_path, service_url, capa_proyecto = data
        
        #print_with_line_number(f"wsdl_path: {wsdl_path}")
        #print_with_line_number(f"operations: {operations}")
        #print_with_line_number(f"project_name: {project_name}")
        #print_with_line_number(f"service_name: {service_name}")
        #print_with_line_number(f"osb_file_path: {osb_file_path}")
        #print_with_line_number(f"pipeline_path: {pipeline_path}")
        #print_with_line_number(f"service_url: {service_url}")
        #print_with_line_number(f"capa_proyecto: {capa_proyecto}")

        imports = extract_xsd_import_paths(wsdl_path)
        #print_with_line_number(f"wsdl_path: {wsdl_path}")
        #print_with_line_number(f"imports: {imports}")
        
        #print_with_line_number(f"project_path: {project_path}")
        # 🔹 Eliminar 'extraccion_jar/' para obtener la ruta relativa base
        wsdl_relative_base = os.path.relpath(wsdl_path, "extraccion_jar")
        #print_with_line_number(f"wsdl_relative_base: {wsdl_relative_base}")
        operacion_business = ""
        # 🔹 Obtener la carpeta donde está el WSDL
        wsdl_dir = os.path.dirname(wsdl_relative_base)
        #print_with_line_number(f"wsdl_dir: {wsdl_dir}")
        # 🔹 Procesar cada import y ajustar solo los que empiezan con "../Schemas"
        xsd_relative_paths = []
        # 🔹 Modificar `imports` en su lugar
        for i, imp in enumerate(imports):
            if imp.startswith("../Schemas"):  # Solo modificar los que empiezan con "../Schemas"
                imports[i] = os.path.normpath(os.path.join(wsdl_dir, imp))  # Reemplazar en la misma lista
                                            
        
        #print_with_line_number(f"imports despues: {imports}")
        
        if operacion_a_documentar in operations or not operacion_a_documentar:
            for operation in operations:
                for xsd in imports:
                    xsd_filename = os.path.basename(xsd).lower()  # Obtener solo el nombre del archivo XSD

                    # 🔹 Buscar coincidencia exacta con el nombre del XSD
                    if xsd_filename == operation.lower() + ".xsd":
                        operation_to_xsd[operation] = xsd
                        break  # Detener la búsqueda cuando encuentra la coincidencia exacta

                else:  # Solo ejecuta este bloque si el `for xsd in imports` no encontró nada
                    xsd_names = [os.path.basename(x).lower() for x in imports]  # Lista de nombres de archivos XSD
                    closest_match = difflib.get_close_matches(operation.lower() + ".xsd", xsd_names, n=1, cutoff=0.9)

                    if closest_match:
                        matched_xsd = next(x for x in imports if os.path.basename(x).lower() == closest_match[0])
                        operation_to_xsd[operation] = matched_xsd
                    else:
                        operation_to_xsd[operation] = None  # No se encontró una coincidencia
            
            #print_with_line_number(f"operation_to_xsd: {operation_to_xsd}")

            # ✅ Si el usuario especificó una operación, verificar si existe en operation_to_xsd
            if operacion_a_documentar and operacion_a_documentar not in operation_to_xsd:
                continue
            else:
                found = True  # La operación se encontró en este archivo
                # Iterar sobre el diccionario y realizar la llamada a parse_xsd_file
                for operation_name, xsd in operation_to_xsd.items():
                    #
                    operation_actual = operation_name
                    #print_with_line_number(f"operation_actual: {operation_actual}")
                    #print_with_line_number(f"operacion_a_documentar: {operacion_a_documentar}")
                    if not operacion_a_documentar or operation_name == operacion_a_documentar:
                        #print_with_line_number(f"operation_actual: {operation_actual}")
                        print_with_line_number(f"🔍 Analizando operacion: {operation_actual}")
                        #print_with_line_number(f"service_name: {service_name}")
                        #print_with_line_number(f"operation_name: {operation_name}")
                        #print_with_line_number(f"service_url: {service_url}")
                        #print_with_line_number(f"capa_proyecto: {capa_proyecto}")
                        #print_with_line_number(f"operacion_business: {operacion_business}")
                        xsd = os.path.splitext(xsd)[0] + ".XMLSchema"
                        #print_with_line_number(f"xsd: {xsd}")
                    
                        elementos_xsd = parse_xsd_file(project_path,xsd, operation_name,service_url,capa_proyecto,operacion_business,operations, service_name, operation_actual)
                        #print_with_line_number(f"elementos_xsd: {elementos_xsd}")

                        #services_for_operations = recorrer_servicios_internos_osb(project_path,operacion_a_documentar,osb_file_path, pipeline_path, operations, visited_proxies)

                        osb_services.append(elementos_xsd)
                    
                        if operacion_a_documentar:
                            return osb_services
                                                    
        if not found:  
            st.error("⛔ No se encuentra la operación en el .jar ⛔")

    #st.success(f"osb_services: {osb_services}")
    return osb_services

def generar_operaciones_expuestas_http(project_path,operacion_a_documentar):
    
    osb_services = []
    elementos_xsd = []
    operations =[]
    operation_to_xsd = {}
    services_for_operations = {}
    combined_services = {}
    found = False  # Variable para rastrear si se encuentra la operación
    
    plantuml_code = """@startuml
    Alice->Bob : I am using hex
    @enduml"""
    
    hex_string = plantuml_to_hex(plantuml_code)
    #print_with_line_number(f"hex_string: {hex_string}")
    
    # URL final
    plantuml_url_png = f"https://www.plantuml.com/plantuml/png/{hex_string}"
    #st.image(plantuml_url_png)
    #print("🔹 URL de la imagen PNG:", plantuml_url_png)
    
    
    uml_example = """@startuml
    Alice -> Bob: Authentication Request
    Bob --> Alice: Authentication Response
    @enduml"""

    # Generar URL
    uml_url = generate_plantuml_url(uml_example)
    #print_with_line_number(f"URL del diagrama: {uml_url}")
    
    # URL final
    plantuml_url_png = {uml_url}
    #st.image(plantuml_url_png)
    #print("🔹 URL de la imagen PNG:", plantuml_url_png)

    ##print_with_line_number(f"URL generada: {url}")
    wsdl_operations_map = extraer_operaciones_expuestas_http(project_path)
    
    # Recorriendo el diccionario
    for wsdl_path, data in wsdl_operations_map.items():
        # Desempaquetar la tupla
        operations, project_name, service_name, osb_file_path, pipeline_path, service_url, capa_proyecto = data
        
        #print_with_line_number(f"wsdl_path: {wsdl_path}")
        #print_with_line_number(f"operations: {operations}")
        #print_with_line_number(f"project_name: {project_name}")
        #print_with_line_number(f"service_name: {service_name}")
        #print_with_line_number(f"osb_file_path: {osb_file_path}")
        #print_with_line_number(f"pipeline_path: {pipeline_path}")
        #print_with_line_number(f"service_url: {service_url}")
        #print_with_line_number(f"capa_proyecto: {capa_proyecto}")

        # for operation in operations:
        if operacion_a_documentar:
            operations = operacion_a_documentar

        services_for_operations_exp = extraer_operaciones_pipeline_exp(pipeline_path, operations)
                    
        #print_with_line_number(f"services_for_operations_exp: {services_for_operations_exp}")
        
        services_for_operations_ebs = extraer_operaciones_pipeline_ebs(project_path,services_for_operations_exp)
        
        #print_with_line_number(f"services_for_operations_ebs: {services_for_operations_ebs}")

        for operation, proxy_list in services_for_operations_exp.items():
            combined_services[operation] = {'Proxy': proxy_list, 'Referencia': []}

        for operation, reference_list in services_for_operations_ebs:
            if operation in combined_services:
                combined_services[operation]['Referencia'] = reference_list
            else:
                combined_services[operation] = {'Proxy': [], 'Referencia': reference_list}
        
        #print_with_line_number(f"combined_services: {combined_services}")
        
        combined_services2 = separar_ebs_abc_business(project_path,combined_services)
        
        #print_with_line_number(f"combined_services2: {combined_services2}")
        
        #generar_diagramas_operaciones(project_name,combined_services2)
            
    return combined_services2

def recorrer_servicios_internos_osb(project_path,operacion_a_documentar,proxy_path, pipeline_path, operations, visited_proxies=None):
    if visited_proxies is None:
        visited_proxies = set()

    services_for_operations = defaultdict(list)
    
    #print_with_line_number(f"🔍 project_path: {project_path}")
    #print_with_line_number(f"🔍 proxy_path: {proxy_path}")
    #print_with_line_number(f"🔍 pipeline_path: {pipeline_path}")

    for operacion_padre in operations:
        operacion_actual = operacion_padre
        buscar_branch_operacion(pipeline_path, project_path, operations, operacion_actual)
        #extract_service_for_operations_audibpel(project_path,pipeline_path,operations,services_for_operations,operacion_padre,operacion_actual)
        #procesar_pipeline(project_path, proxy_path,pipeline_path, operacion_padre)
    
    #print_with_line_number(f"Servicios internos encontrados: {services_for_operations}")
    return services_for_operations

def procesar_pipeline(project_path, proxy_actual, pipeline_actual, operacion_actual=None, services_for_operations=None):
    if services_for_operations is None:
        services_for_operations = defaultdict(dict)

    namespaces = {
        'con': 'http://www.bea.com/wli/sb/pipeline/config',
        'con1': 'http://www.bea.com/wli/sb/stages/routing/config',
        'con2': 'http://www.bea.com/wli/sb/stages/config',
        'con3': 'http://www.bea.com/wli/sb/stages/transform/config',
        'con4': 'http://www.bea.com/wli/sb/stages/publish/config',
        'ref': 'http://www.bea.com/wli/sb/reference',
        'xsi': 'http://www.w3.org/2001/XMLSchema-instance'
    }

    #print_with_line_number(f"🔍 project_path: {project_path}")
    #print_with_line_number(f"🔍 proxy_actual: {proxy_actual}")
    #print_with_line_number(f"🔍 pipeline_actual: {pipeline_actual}")

    if not os.path.exists(pipeline_actual):
        st.warning(f"Archivo no encontrado: {pipeline_actual}")
        return services_for_operations

    with open(pipeline_actual, "r", encoding="utf-8") as file:
        xml_content = file.read()
    root = ET.fromstring(xml_content)

    # Obtener operaciones del WSDL asociado al pipeline
    wsdl_pipeline = ""
    wsdl_element = root.find('.//con:wsdl', namespaces)
    if wsdl_element is not None and 'ref' in wsdl_element.attrib:
        wsdl_pipeline = os.path.join(project_path, wsdl_element.attrib['ref'] + ".WSDL")
        st.info(f"wsdl_pipeline: {wsdl_pipeline}")
        operations = extract_wsdl_operations(wsdl_pipeline)
        st.info(f"operations: {operations}")

        # Iterar sobre cada operación principal del pipeline
        for operacion_padre in operations:
            #print_with_line_number(f"🔍 operacion_padre: {operacion_padre}")

            # Diccionario para registrar servicios invocados en esta operación
            referencias = []

            # Buscar en `branch` referencias a otros ProxyService
            branch_xpath = f".//con:branch[@name='{operacion_padre}']"
            branch = root.find(branch_xpath, namespaces)
            if branch is not None:
                for service in branch.findall(".//con1:service[@xsi:type='ref:ProxyRef']", namespaces):
                    service_ref = service.get("ref")
                    #print_with_line_number(f"🔍1 service_ref: {service_ref}")
                    if service_ref:
                        initial_proxy_path = os.path.join(project_path, service_ref + ".ProxyService")
                        #print_with_line_number(f"🔍1 initial_proxy_path: {initial_proxy_path}")
                        new_pipeline_path = extract_pipeline_path_from_proxy(initial_proxy_path, project_path)
                        #print_with_line_number(f"🔍1 new_pipeline_path: {new_pipeline_path}")

                        # Recursivamente procesar el pipeline hijo
                        sub_operations = procesar_pipeline(
                            project_path, initial_proxy_path, new_pipeline_path, operacion_padre, services_for_operations
                        )

                        referencias.append((service_ref, sub_operations))

            # Buscar en `routes` referencias a BusinessService
            routes = root.findall(".//con1:route", namespaces)
            for route in routes:
                business_service = route.find(".//con1:service", namespaces)
                operation = route.find(".//con1:operation", namespaces)

                if business_service is not None and "ref" in business_service.attrib:
                    service_ref = business_service.attrib["ref"]
                    #print_with_line_number(f"🔍2 service_ref: {service_ref}")
                    operation_name = operation.text if operation is not None else ""
                    referencias.append((service_ref, operation_name))
                    #print_with_line_number(f"BusinessService detectado: {service_ref} con operación {operation_name}")

            # Almacenar las referencias de la operación padre en el diccionario principal
            if referencias:
                services_for_operations[operacion_padre][pipeline_actual] = referencias
                #print_with_line_number(f"🔍 services_for_operations actualizado: {services_for_operations}")

    return services_for_operations

def buscar_branch_operacion(pipeline_path, project_path, operations, operacion_a_documentar):
    if pipeline_path.endswith('.Pipeline') and os.path.isfile(pipeline_path):
        #print_with_line_number(f"📂 Analizando pipeline: {pipeline_path}")

        # Leer el contenido del pipeline
        with open(pipeline_path, 'r', encoding="utf-8") as f:
            pipeline_content = f.read()
        
        # Cargar el XML
        root = ET.fromstring(pipeline_content)
        
        # Definir los namespaces del XML
        namespaces = {
            'con': 'http://www.bea.com/wli/sb/pipeline/config', 
            'con1': 'http://www.bea.com/wli/sb/stages/routing/config',
            'con2': 'http://www.bea.com/wli/sb/stages/config',
            'con3': 'http://www.bea.com/wli/sb/stages/transform/config',
            'con4': 'http://www.bea.com/wli/sb/stages/publish/config',
            'ref': 'http://www.bea.com/wli/sb/reference',
            'xsi': 'http://www.w3.org/2001/XMLSchema-instance'
        }
        
        lista_proxys = []
        lista_operaciones_proxys = []
        
        # Buscar el <con:branch> con el name específico
        branch_xpath = f".//con:branch[@name='{operacion_a_documentar}']"
        branch_element = root.find(branch_xpath, namespaces)
        
        if branch_element is not None:
            #print_with_line_number(f"✅ Se encontró el branch: {operacion_a_documentar}")
            
            # Buscar el <con1:service> dentro del branch encontrado
            service_element = branch_element.find(".//con1:service", namespaces)
            
            if service_element is not None:
                service_ref = service_element.attrib.get('ref', '')
                #print_with_line_number(f"🔗 Referencia al servicio: {service_ref}")
                
                # Construir la ruta al ProxyService
                proxy_referencia = os.path.join(project_path, service_ref + ".ProxyService")
                #print_with_line_number(f"📄 Proxy referencia: {proxy_referencia}")
                
                # Obtener el pipeline asociado al proxy
                new_pipeline_path = extract_pipeline_path_from_proxy(proxy_referencia, project_path)
                #print_with_line_number(f"📂 Nuevo pipeline detectado: {new_pipeline_path}")

                return new_pipeline_path
            #else:
                #print_with_line_number("⚠️ No se encontró un <con1:service> dentro del branch.")
        #else:
            #print_with_line_number(f"❌ No se encontró el branch con name='{operacion_a_documentar}' en el pipeline.")
        
    return None

def extraer_operaciones_pipeline_exp(pipeline_path, operations):
    services_for_operations = defaultdict(set)
    
    #print_with_line_number("***************************** INICIO EXTRACT SERVICE OPERATIONS*********************************************")

    if not (pipeline_path.endswith('.Pipeline') and os.path.isfile(pipeline_path)):
        #print_with_line_number("Archivo no válido o no encontrado.")
        return services_for_operations

    #print_with_line_number(f"pipeline_path: {pipeline_path}")

    # Cargar el archivo XML
    with open(pipeline_path, 'r', encoding="utf-8") as f:
        root = ET.fromstring(f.read())

    namespaces = {
        'con': 'http://www.bea.com/wli/sb/pipeline/config',
        'con1': 'http://www.bea.com/wli/sb/stages/routing/config',
        'con2': 'http://www.bea.com/wli/sb/stages/config',
        'con3': 'http://www.bea.com/wli/sb/stages/transform/config',
        'con4': 'http://www.bea.com/wli/sb/stages/publish/config',
        'ref': 'http://www.bea.com/wli/sb/reference',
        'xsi': 'http://www.w3.org/2001/XMLSchema-instance'
    }

    def process_branch_elements():
        """Busca servicios en elementos <con:branch>."""
        for branch in root.findall(".//con:branch", namespaces):
            operation_name = branch.attrib.get('name', '')
            if operation_name in operations:
                service_element = branch.find(".//con1:service", namespaces)
                if service_element is not None:
                    services_for_operations[operation_name].add(service_element.attrib.get('ref', ''))
                    #print_with_line_number(f"services_for_operations process_branch_elements: {services_for_operations}")
        return services_for_operations

    def process_flow_elements():
        """Busca servicios en elementos <con:flow>."""
        for flow in root.findall(".//con:flow", namespaces):
            for service_element in flow.findall(".//con1:service[@xsi:type='ref:BusinessServiceRef']", namespaces):
                service_ref = service_element.attrib.get('ref', '')
                for operation_element in flow.findall(".//con1:operation", namespaces):
                    operation_name = operation_element.text.strip()
                    if operation_name in operations:
                        services_for_operations[operation_name].add(service_ref)
                        #print_with_line_number(f"services_for_operations process_flow_elements: {services_for_operations}")
        return services_for_operations

    def process_route_elements():
        """Busca servicios en elementos <con:route-node>."""
        for route in root.findall(".//con:route-node", namespaces):
            operation_element = route.find(".//con1:operation", namespaces)
            if operation_element is not None:
                operation_name = operation_element.text.strip()
                if operation_name in operations:
                    service_element = route.find(".//con1:service", namespaces)
                    if service_element is not None:
                        services_for_operations[operation_name].add(service_element.attrib.get('ref', ''))
                        #print_with_line_number(f"services_for_operations process_route_elements: {services_for_operations}")
        return services_for_operations

    def process_callout_elements():
        """Busca servicios en elementos <wsCallout>."""
        for callout in (e for e in root.iter() if e.tag.endswith('wsCallout')):
            operation_element = callout.find(".//con3:operation", namespaces)
            service_element = callout.find(".//con3:service", namespaces)
            if operation_element is not None and service_element is not None:
                operation_name = operation_element.text.strip()
                if operation_name in operations:
                    services_for_operations[operation_name].add(service_element.attrib.get('ref', ''))
                    #print_with_line_number(f"services_for_operations process_callout_elements: {services_for_operations}")
        return services_for_operations

    
    branch_found = process_branch_elements()
    flow_found = process_flow_elements()
    route_found = process_route_elements()
    callout_found = process_callout_elements()
    
    # Ejecutar los procesamientos en orden hasta encontrar un servicio
    seguir = True

    #print_with_line_number(f"SERVICES FOR: {dict(services_for_operations)}")
    #print_with_line_number("***************************** FIN EXTRACT SERVICE OPERATIONS*********************************************")

    return {op: list(set(services)) for op, services in services_for_operations.items()}

def extraer_operaciones_pipeline_ebs(jdeveloper_projects_dir, services_for_operations):
    osb_services = []
    
    for operacion, paths in services_for_operations.items():
        for path2 in paths:
            #print_with_line_number("********** INICIO PROCESO **********")
            #print_with_line_number(f"Operacion: {operacion}, Path: {path2}")
            
            if 'Proxies' in path2:
                osb_file_path = os.path.join(jdeveloper_projects_dir, path2 + ".ProxyService")
            elif 'Pipeline' in path2:
                osb_file_path = os.path.join(jdeveloper_projects_dir, path2 + ".Pipeline")
            else:
                continue
            
            #print_with_line_number(f"OSB File Path: {osb_file_path}")
            
            project_name = extract_project_name_from_proxy(osb_file_path)
            if project_name is None:
                continue
            
            pipeline_path = osb_file_path if 'Pipeline' in path2 else extract_pipeline_path_from_proxy(osb_file_path, jdeveloper_projects_dir)
            #print_with_line_number(f"Pipeline Path: {pipeline_path}")
            
            with open(osb_file_path, 'r', encoding="utf-8") as f:
                content = f.read()
                service_name = os.path.splitext(os.path.basename(osb_file_path))[0]
                wsdl_relative_path = extract_wsdl_relative_path(content)
            
            #print_with_line_number(f"Service Name: {service_name}")
            #print_with_line_number(f"WSDL Relative Path: {wsdl_relative_path}")
            
            if wsdl_relative_path:
                wsdl_path = os.path.join(jdeveloper_projects_dir, wsdl_relative_path + ".WSDL")
                operations = extract_wsdl_operations(wsdl_path)
                #print_with_line_number(f"Operations: {operations}")
            
            #print_with_line_number(f"Pipeline Path: {pipeline_path}")
            service_for_operations = definir_operaciones_internas_pipeline(pipeline_path)
            #print_with_line_number(f"Service for Operations: {service_for_operations}")
            
            if service_for_operations:
                rutas_de_servicio = list(service_for_operations.values())
                osb_services.append((operacion, rutas_de_servicio))
            #else:
                #service_refs = extract_service_refs_from_pipeline(pipeline_path)
                #osb_services.append((operacion, path2))
                #print_with_line_number(f"Service Refs: {service_refs}")
    
    #print_with_line_number("********** FIN PROCESO **********")
    return osb_services

def definir_operaciones_internas_pipeline(pipeline_path):
    service_refs = set()
    services_for_operations = {}
    
    #print_with_line_number(f"pipeline_path: {pipeline_path}")
    
    namespaces = {
        'transform': 'http://www.bea.com/wli/sb/stages/transform/config',
        'publish': 'http://www.bea.com/wli/sb/stages/publish/config',
        'routing': 'http://www.bea.com/wli/sb/stages/routing/config',
        'config': 'http://www.bea.com/wli/sb/stages/config',
        'pipeline': 'http://www.bea.com/wli/sb/pipeline/config',
        'ref': 'http://www.bea.com/wli/sb/reference',
        'xsi': 'http://www.w3.org/2001/XMLSchema-instance'
    }
    
    try:
        with open(pipeline_path, 'r', encoding="utf-8") as f:
            root = ET.fromstring(f.read())

        def extract_services_and_operations(elements, ns, service_tag, operation_tag):
            for element in elements:
                service_element = element.find(f".//{ns}:{service_tag}", namespaces)
                operation_element = element.find(f".//{ns}:{operation_tag}", namespaces)
                if service_element is not None and operation_element is not None:
                    service_ref = service_element.attrib.get('ref', '')
                    operation_name = operation_element.text.strip()
                    service_refs.add(service_ref)
                    
                    new_operation_name = operation_name
                    version = 2
                    while new_operation_name in services_for_operations and services_for_operations[new_operation_name] != service_ref:
                        new_operation_name = f"{operation_name}v{version}"
                        version += 1
                    
                    services_for_operations[new_operation_name] = service_ref

        extract_services_and_operations(root.findall(".//transform:wsCallout", namespaces), 'transform', 'service', 'operation')
        extract_services_and_operations(root.findall(".//config:wsCallout", namespaces), 'config', 'service', 'operation')
        extract_services_and_operations(root.findall(".//publish:route", namespaces), 'publish', 'service', 'operation')
        extract_services_and_operations(root.findall(".//routing:route", namespaces), 'routing', 'service', 'operation')
        extract_services_and_operations(root.findall(".//pipeline:flow", namespaces), 'pipeline', 'service', 'operation')
        
        return services_for_operations
    
    except Exception as e:
        print(f"Error procesando el pipeline: {e}")
        return {}

def separar_ebs_abc_business(jdeveloper_projects_dir, combined_services):
    """
    Recorre recursivamente las referencias de un servicio en busca de dependencias,
    actualizando `combined_services` con referencias y detalles de BusinessServices.
    """
    
    informacion_business = {}
    referencias = {}
    referencia_proxy ={}
    referencia_business_service ={}
    
    def buscar_recursivamente_operaciones(referencia):
        referencia_base = os.path.basename(referencia)  # Extrae solo el nombre del archivo
        referencia_base = referencia_base.replace(".ProxyService", "").replace(".BusinessService", "")  # Normaliza nombres
        #print_with_line_number(f"🔍 referencia_base: {referencia_base}")
        if "Proxies" in referencia:
            osb_file_path = os.path.join(jdeveloper_projects_dir, referencia + ".ProxyService")
            if os.path.exists(osb_file_path):
                #print_with_line_number(f"🔍 osb_file_path: {osb_file_path}")
                project_name = extract_project_name_from_proxy(osb_file_path)
                #print_with_line_number(f"🔍 project_name: {project_name}")
                pipeline_path = extract_pipeline_path_from_proxy(osb_file_path, jdeveloper_projects_dir)
                #print_with_line_number(f"🔍 pipeline_path: {pipeline_path}")
                service_for_operations = definir_operaciones_internas_pipeline(pipeline_path)
                #print_with_line_number(f"🔍 service_for_operations: {service_for_operations}")

                if service_for_operations:
                    referencias[f"REFERENCIA_{referencia_base}"] = service_for_operations
                    for valor in service_for_operations.values():
                        valor_buscado = valor
                        #print_with_line_number(f"🔍 valor_buscado: {valor_buscado}")

                        if "BusinessServices" in valor_buscado:
                            referencia_business_service = valor_buscado
                            #print_with_line_number(f"🔍 referencia_business_service: {referencia_business_service}")
                            biz_path = os.path.join(jdeveloper_projects_dir, referencia_business_service + ".BusinessService")
                            #print_with_line_number(f"🔍 biz_path: {biz_path}")
                            if os.path.exists(biz_path):
                                service_refs = extract_uri_and_provider_id_from_bix(biz_path)
                                if service_refs:
                                    informacion_business[f"INFORMACION_{referencia_business_service}"] = service_refs
                                    #return informacion_business
                    
                        elif "Proxies" in valor_buscado:
                            referencia_proxy = valor_buscado
                            #print_with_line_number(f"🔍 referencia_proxy: {referencia_proxy}")
                            buscar_recursivamente_operaciones(referencia_proxy)

        elif "BusinessServices" in referencia:
            biz_path = os.path.join(jdeveloper_projects_dir, referencia + ".BusinessService")
            #print_with_line_number(f"🔍 biz_path: {biz_path}")
            if os.path.exists(biz_path):
                service_refs = extract_uri_and_provider_id_from_bix(biz_path)
                if service_refs:
                    informacion_business[f"INFORMACION_{referencia_base}"] = service_refs
                    #return informacion_business
        
        
        elif "Pipeline" in referencia:
            pipeline_path = os.path.join(jdeveloper_projects_dir, referencia + ".Pipeline")
            #print_with_line_number(f"🔍 pipeline_path: {pipeline_path}")
            service_for_operations = definir_operaciones_internas_pipeline(pipeline_path)
            #print_with_line_number(f"🔍 service_for_operations: {service_for_operations}")

            if service_for_operations:
                referencias[f"REFERENCIA_{referencia_base}"] = service_for_operations
                for valor in service_for_operations.values():
                    valor_buscado = valor
                    #print_with_line_number(f"🔍 valor_buscado: {valor_buscado}")

                    if "BusinessServices" in valor_buscado:
                        referencia_business_service = valor_buscado
                        #print_with_line_number(f"🔍 referencia_business_service: {referencia_business_service}")
                        biz_path = os.path.join(jdeveloper_projects_dir, referencia_business_service + ".BusinessService")
                        #print_with_line_number(f"🔍 biz_path: {biz_path}")
                        if os.path.exists(biz_path):
                            service_refs = extract_uri_and_provider_id_from_bix(biz_path)
                            if service_refs:
                                informacion_business[f"INFORMACION_{referencia_business_service}"] = service_refs
                                #return informacion_business
                
                    elif "Proxies" in valor_buscado:
                        referencia_proxy = valor_buscado
                        #print_with_line_number(f"🔍 referencia_proxy: {referencia_proxy}")
                        buscar_recursivamente_operaciones(referencia_proxy)
    
    
    for service_name, service_data in combined_services.items():
        #print_with_line_number(f"🔍 servicio: {service_name}")
        for proxy in service_data.get("Proxy", []):
            for referencia in service_data.get("Referencia", []):
                if proxy not in referencia:
                    operacion = buscar_recursivamente_operaciones(referencia)
                
        # Actualizar el servicio actual en combined_services con la nueva información
        combined_services[service_name].update(referencias)
        combined_services[service_name].update(informacion_business)
        #print_with_line_number(f"🔍 combined_services[service_name]: {combined_services[service_name]}")
        informacion_business = {}
        referencias = {}
        referencia_proxy ={}
        referencia_business_service ={}
    
    return combined_services
    
def separar_ebs_abc_business2(jdeveloper_projects_dir,combined_services):
    
    for service_name, service_data in combined_services.items():
        informacion_business = {}
        referencias = {}
        #print_with_line_number(f"service_name: {service_name}")
        
        #for proxy in service_data.get("Proxy", []):
            #print_with_line_number(f"proxy: {proxy}")

        # Recorrer las Referencias
        for referencia in service_data.get("Referencia", []):
            #print_with_line_number(f"referencia: {referencia}")
            
            if "Proxies" in referencia:
                osb_file_path = os.path.join(jdeveloper_projects_dir, referencia + ".ProxyService")
                if os.path.exists(osb_file_path):
                    #print_with_line_number(f"🔍osb_file_path: {osb_file_path}")
                    project_name = extract_project_name_from_proxy(osb_file_path)
                    #print_with_line_number(f"🔍project_name: {project_name}")
                    pipeline_path = extract_pipeline_path_from_proxy(osb_file_path, jdeveloper_projects_dir)
                    #print_with_line_number(f"🔍pipeline_path: {pipeline_path}")
                    service_for_operations = definir_operaciones_internas_pipeline(pipeline_path)
                    #print_with_line_number(f"🔍service_for_operations: {service_for_operations}")
                
                    if service_for_operations:
                        rutas_de_servicio = list(service_for_operations.values())
                        referencias[f"REFERENCIA_{os.path.basename(referencia)}"] = rutas_de_servicio
                        
            
            if "BusinessServices" in referencia:
                biz_path = os.path.join(jdeveloper_projects_dir, referencia + ".BusinessService")
                
                if os.path.exists(biz_path):
                    service_refs = extract_uri_and_provider_id_from_bix(biz_path)
                    if service_refs:
                        informacion_business[f"INFORMACION_{os.path.basename(referencia)}"] = service_refs
        
        service_data.update(referencias)
        service_data.update(informacion_business)
        #print_with_line_number(f"service_data: {service_data}")
    #print_with_line_number(f"combined_services: {combined_services}")
    
def extract_uri_and_provider_id_from_bix(bix_path):
    lista_uri_provider = []
    with open(bix_path, 'r', encoding="utf-8") as f:
        content = f.read()
        # Buscar el valor dentro de las etiquetas <env:value>
        uri_match = re.search(r'<env:value>(.*?)</env:value>', content, re.DOTALL)
        
        #print_with_line_number(f"MATCH: {uri_match}")
        if uri_match:
            uri_value = uri_match.group(1)
        else:
            uri_value = None

        #print_with_line_number(f"URI VALUE: {uri_value}")
        # Buscar el valor dentro de las etiquetas <tran:provider-id>
        provider_id_match = re.search(r'<tran:provider-id>(.*?)</tran:provider-id>', content, re.DOTALL)
        #print_with_line_number(f"PROVIDER_ID: {provider_id_match}")
        if provider_id_match:
            provider_id_value = provider_id_match.group(1)
        else:
            provider_id_value = None
        
        #print_with_line_number(f"PROVIDER_ID_VALUE: {provider_id_value}")
        lista_uri_provider.append((uri_value, provider_id_value))
        return lista_uri_provider

def reemplazar_marcador_con_imagen(doc, marcador, diagrama_path):
    """
    Busca un marcador en el documento y lo reemplaza con una imagen en una página completa en orientación horizontal.
    Retorna el documento modificado.
    """
    for para in doc.paragraphs:
        if marcador in para.text:
            # Agregar una nueva sección con orientación horizontal
            section = para._element.getparent().addnext(doc.add_section()._element)
            new_section = doc.sections[-1]
            new_section.orientation = WD_ORIENT.LANDSCAPE
            new_section.page_width = Cm(29.7)  # A4 Horizontal
            new_section.page_height = Cm(21.0)
            new_section.left_margin = Cm(1.0)
            new_section.right_margin = Cm(1.0)
            new_section.top_margin = Cm(1.0)
            new_section.bottom_margin = Cm(1.0)

            # Limpiar el marcador
            para.text = para.text.replace(marcador, "")

            # Insertar imagen en la nueva sección
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            if os.path.exists(diagrama_path):
                run.add_picture(diagrama_path, width=Cm(27))  # Ajusta el tamaño de la imagen
            else:
                print(f"ERROR: No se encontró la imagen {diagrama_path}")

            return doc  # Retornar el documento modificado
    return doc  # Retornar el documento si no se encontró el marcador


def generar_documentacion(jar_path, plantilla_path,operacion_a_documentar,nombre_autor):
    """Función que ejecuta la generación de documentación."""
    
    zip_files = []
    generoArchivo = False
    
    # Extraer ruta del proyecto desde el .jar
    jdeveloper_projects_dir = jar_path
    
    #print_with_line_number(f"✅ jdeveloper_projects_dir {jdeveloper_projects_dir}")
    
    if not jdeveloper_projects_dir:
        st.error("No se pudo determinar la ruta del proyecto desde el .jar.")
        return

    # 📌 Definir la ruta del directorio temporal correctamente
    temp_dir = os.path.join(tempfile.gettempdir(), "documentacion_osb")
    ruta_temporal = temp_dir  # Obtener la ruta temporal

    if not isinstance(temp_dir, str) or not temp_dir:
        st.error("⛔ Error: La ruta temporal no es válida.")
    else:
        # 📌 Verificar si la carpeta existe antes de intentar eliminarla
        if os.path.exists(temp_dir):
            try:
                shutil.rmtree(temp_dir)  # 🔥 Borra todo el contenido anterior
                #st.warning("📂 Se limpiaron los archivos temporales previos.")
            except Exception as e:
                st.error(f"⛔ No se pudo eliminar la carpeta temporal: {e}")

        # 📌 Crear nuevamente la carpeta temporal limpia
        os.makedirs(temp_dir, exist_ok=True)
        #print_with_line_number(f"📂 Carpeta temporal creada: {temp_dir}")
    
    # Llamar a la función principal de tu script
    services_with_data = extraer_schemas_operaciones_expuestas_http(jdeveloper_projects_dir,operacion_a_documentar)
    
    #print_with_line_number(f"✅ services_with_data {services_with_data}")
    
    es_type = False
    
    # Initialize an empty set to store unique operation names
    operation_names = set()
    
    if services_with_data:

        # Iterate through each tuple of request and response elements in services_with_data
        for request_elements, response_elements in services_with_data:
            # Iterate through each element in request_elements and response_elements
            for element in request_elements + response_elements:
                if 'Type' in element['elemento']:
                    es_type = True
                #operation_name = element['elemento'].replace('Request', '').replace('Response', '').replace('Type', '')
                ##print_with_line_number(f"operation_name: {operation_name}")
                service_name = element['service_name']
                # Agregar todas las operaciones de la lista 'operations'
                if 'operations' in element:
                    operation_names.update(element['operations'])  # Agrega todas las operaciones a operation_names

        # Convert the set to a sorted list to get the operation names in alphabetical order
        unique_operations = sorted(operation_names)
        
        operaciones_formateadas = "\n".join(f"* {op}" for op in unique_operations)
        
        
        # 🔹 Si operacion_a_documentar tiene un valor, filtrar solo esa operación
        if operacion_a_documentar:
            unique_operations = [operacion_a_documentar] if operacion_a_documentar in unique_operations else []
            
        
        #print_with_line_number(f"unique_operations: {unique_operations}")
        
        #print_with_line_number(f"✅ unique_operations {unique_operations}")
        
        operation_elements = {}
        
        
        total_operaciones = len(unique_operations)
        if total_operaciones == 0:
            st.warning("⚠️ No hay operaciones que documentar.")
            return
        
        if total_operaciones > 1:
            progress_bar_general = st.progress(0)
        
        # 🔹 Iterar sobre cada operación
        for idx, operation in enumerate(unique_operations, start=1):
            if total_operaciones > 1:
                progreso_actual = int((idx / total_operaciones) * 100)
                progress_bar_general.progress(progreso_actual)  # 🔄 Actualizar barra general
                #print_with_line_number(f"⏳ Procesando operación {idx}/{total_operaciones}: {operation} ({progreso_actual}%)")
            else:
                print_with_line_number(f"⏳ Procesando operación {idx}/{total_operaciones}: {operation}")
            
            
            if es_type:
                request_key = f"{operation}RequestType"
                response_key = f"{operation}ResponseType"
            else:
                request_key = f"{operation}Request"
                response_key = f"{operation}Response"
            
            # Initialize lists to store request and response elements for the current operation
            request_elements = []
            response_elements = []
            url_elements = []
            capa_proyecto = []
            minOccurs_elements = []
            
            # Iterate through services_with_data to find matching elements
            for request_data, response_data in services_with_data:
                #print_with_line_number(f"request_data: {request_data}")
                # Check for request elements
                for element in request_data:
                    elemento_nombre = element['elemento']
                    # ✅ Verificar coincidencia exacta o parcial usando difflib
                    match = difflib.get_close_matches(request_key, [elemento_nombre], n=1, cutoff=0.9)
                    
                    if match or request_key in elemento_nombre:  # Si hay coincidencia razonable
                        request_elements.append({'name': element['name'], 'type': element['type'],'minOccurs': element['minOccurs']})
                        url_elements.append({'url': element['url']})
                        capa_proyecto.append({'ruta': element['ruta']})
                        minOccurs_elements.append({'minOccurs': element['minOccurs']})
                        service_name = element['service_name']
                
                # 🔹 Verificar si `response_key` está en `response_data['elemento']`
                for element in response_data:
                    elemento_nombre = element['elemento']

                    # ✅ Verificar coincidencia exacta o parcial
                    match = difflib.get_close_matches(response_key, [elemento_nombre], n=1, cutoff=0.9)
                    
                    if match or response_key in elemento_nombre:  
                        response_elements.append({'name': element['name'], 'type': element['type'],'minOccurs': element['minOccurs']})
                        service_name = element['service_name']
            
            # Store the collected elements in the dictionary
            operation_elements[operation] = {
                'request': request_elements,
                'response': response_elements,
                'url': url_elements,
                'ruta': capa_proyecto, 
                'minOccurs': minOccurs_elements,
                'service_name': service_name
            }
        #print_with_line_number(f"operation_elements: {operation_elements}")
        ##print_with_line_number(f"service_name: {service_name}")
        # Print the result
        # 📂 Crear un solo ZIP para todas las operaciones
        zip_buffer = tempfile.NamedTemporaryFile(delete=False, suffix=".zip")
        zip_path = zip_buffer.name  # Ruta del archivo ZIP
        
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
            for idx, (operation, elements) in enumerate(operation_elements.items(), start=1):
                
                #st.write(f"🔹 Procesando operación: {operation}")
                st.write(f"📌 Cantidad de elementos request: {len(elements['request'])}")
                st.write(f"📌 Cantidad de elementos response: {len(elements['response'])}")

                #print_with_line_number(f"elements['request']: {elements['request']}")
                if not elements['request']:
                    st.warning(f"⚠️ La operación {operation} no tiene elementos de entrada, saltando...")
                    continue  # Si no hay request, no genera el documento

                # 🔹 Actualizar progreso de generación de documentos
                if total_operaciones > 1:
                    progreso_actual = int(((idx + total_operaciones) / (total_operaciones * 2)) * 100)
                    progress_bar_general.progress(progreso_actual)

                if elements['request']:
                    
                    st.write(f"🔹 Proyecto {elements['ruta'][0]['ruta'].lstrip('/')}")
                    st.write(f"⏳ Creando documentacion operacion: {operation}")
                    
                    #if total_operaciones == 1:
                        #progress_bar_general = st.progress(2)
                    
                    contiene_cabecera_entrada = False
                    contiene_cabecera_salida = False
                    
                    if any('cabeceraEntrada.' in elem['name'] for elem in elements['request']):
                        #st.write("Se encontró al menos un elemento con '.cabeceraEntrada.'")
                        contiene_cabecera_entrada = True
                    
                    if any('cabeceraSalida.' in elem['name'] for elem in elements['response']):
                        #st.write("Se encontró al menos un elemento con '.cabeceraSalida.'")
                        contiene_cabecera_salida = True
                        
                    # Cargar el documento de la plantilla
                    doc = Document(plantilla_path)
                    
                    # Contar el número de tablas en el documento
                    num_tables = len(doc.tables)
                    
                    #print_with_line_number(f"El documento contiene {num_tables} tabla(s).")

                    # Mostrar cada tabla
                    # for i, table in enumerate(doc.tables):
                        # #print_with_line_number(f"\nTabla {i+1}:")
                        # for row in table.rows:
                            # row_data = [cell.text for cell in row.cells]
                            # print_with_line_number('\t'.join(row_data))
                    
                    url = ""
                    ruta =""
                    minOccurs = ""
                    
                    for elem in elements['url']:
                        url = elem['url']
                        
                    for elem in elements['ruta']:
                        ruta = elem['ruta']
                    
                    for elem in elements['minOccurs']:
                        minOccurs = elem['minOccurs']
                        
                    #st.success(f"url: {url}")
                    
                    #st.success(f"ruta: {ruta}")
                    
                    #st.success(f"business: {business}")
                    
                    fecha_actual = datetime.now()
                    fecha_formateada = fecha_actual.strftime("%d/%m/%Y")
                    
                    
                    ruta_proyecto = ruta.strip("/") 
                    
                    combined_services = generar_operaciones_expuestas_http(jdeveloper_projects_dir,operacion_a_documentar)
                    
                    print_with_line_number(f"combined_services: {combined_services}")
                    
                    #print_with_line_number(f"operation: {operation}")
                    
                    diagrama_path = generar_diagramas_operaciones(ruta_proyecto,service_name, combined_services, operation)
                    
                    print_with_line_number(f"diagrama_path: {diagrama_path}")
                    
                    if os.path.exists(diagrama_path):
                        #doc = reemplazar_marcador_con_imagen(doc, "{Imagen_diagrama}", diagrama_path)
                        marcador = "{Imagen_diagrama}"
                        # Obtener el ancho de la página disponible
                        section = doc.sections[0]  # Suponemos que la plantilla tiene una sola sección horizontal
                        #print_with_line_number(f"section: {section}")
                        page_width = section.page_width
                        left_margin = section.left_margin
                        right_margin = section.right_margin

                        # Calcular el ancho disponible para la imagen
                        max_width = page_width - left_margin - right_margin

                        for para in doc.paragraphs:
                            if marcador in para.text:
                                #print_with_line_number(f"Insertando imagen en el marcador: {marcador}")
                                para.text = para.text.replace(marcador, "")  # Borrar el texto del marcador
                                run = para.add_run()
                                run.add_picture(diagrama_path, width=max_width)  # Ajustar la imagen al ancho máximo
                                break  # Solo reemplazamos la primera coincidencia
                    
                    
                    #st.success(f"operation: {operation}")
                    
                    #st.success(f"elements: {elements}")
                    
                    
                    
                    # Definir las variables y sus valores
                    variables = {
                        '{nombre_servicio_inicial}': service_name,
                        '{nombre_servicio_secundario}': service_name,
                        '{nombre_servicio}': service_name,
                        '{nombre_operacion_inicial}' : operation,
                        '{nombre_operacion}': operation,
                        '{unique_operations}': operaciones_formateadas,
                        '{nombre_servicio_contrato}': service_name,
                        '{nombre_servicio_wsdl}': service_name,
                        '{nombre_servicio_contrato2}': service_name,
                        '{nombre_servicio_tabla}': operation,
                        '{fecha}': fecha_formateada,
                        '{autor_inicial}': nombre_autor,
                        '{autor}': nombre_autor,
                        '{autor2}': 'Julian Orjuela',
                        '{url}': url,
                        '{operacion_legado}': minOccurs,
                        '{proyecto_abc}': 'TENENCIA_COMPORTAMIENTO_ABC'
                        # Añade más variables según sea necesario
                    }
                    #st.success(f"service_name: {service_name}")
                    #st.success(f"variables: {variables}")
                    
                    total_tablas = len(doc.tables)
                    #st.success(f"🔍 Total de tablas en el documento: {total_tablas}")
                    if total_operaciones == 1:
                        progress_bar_general = st.progress(30)
                    
                    tabla_cabecera_entrada_numero = 4
                    tabla_cabecera_entrada = doc.tables[tabla_cabecera_entrada_numero - 1]  # Las tablas se indexan desde 0, por eso restamos 1

                    tabla_request_numero = 5
                    tabla_request = doc.tables[tabla_request_numero - 1]  # Las tablas se indexan desde 0, por eso restamos 1
                    
                    tabla_cabecera_salida_numero = 6
                    tabla_cabecera_salida = doc.tables[tabla_cabecera_salida_numero - 1]  # Las tablas se indexan desde 0, por eso restamos 1
                    
                    tabla_response_numero = 7
                    tabla_response = doc.tables[tabla_response_numero - 1]  # Las tablas se indexan desde 0, por eso restamos 1
                    
                    if tabla_cabecera_salida_numero > total_tablas:
                        st.error(f"⛔ Error: Se intentó acceder a la tabla {tabla_cabecera_salida_numero}, pero el documento solo tiene {total_tablas} tablas.")
                        return  # Salir para evitar el error
                    
                    # Listas para almacenar las filas de cada subtabla
                    cabecera_salida = []
                    datos_respuesta = []
                    
                    # Variables de control
                    seccion_actual = None
                    
                    # Datos por defecto para LONGITUD y OBSERVACIÓN
                    default_longitud = "default"
                    default_observacion = ""
                    
                    # Limpiar la tabla antes de agregar elementos de esta operación
                    if not contiene_cabecera_entrada:
                        tbl = tabla_cabecera_entrada._element
                        tbl.getparent().remove(tbl)
                        while len(tabla_cabecera_entrada.rows) > 1:
                            tabla_cabecera_entrada._element.remove(tabla_cabecera_entrada.rows[1]._element)
                            
                    # Limpiar la tabla antes de agregar elementos de esta operación
                    if not contiene_cabecera_salida:
                        tbl = tabla_cabecera_salida._element
                        tbl.getparent().remove(tbl)
                        while len(tabla_cabecera_salida.rows) > 1:
                            tabla_cabecera_salida._element.remove(tabla_cabecera_salida.rows[1]._element)
                    
                    # Limpiar la tabla antes de agregar elementos de esta operación
                    while len(tabla_cabecera_entrada.rows) > 2:
                        tabla_cabecera_entrada._element.remove(tabla_cabecera_entrada.rows[2]._element)
                        
                    # Limpiar la tabla antes de agregar elementos de esta operación
                    while len(tabla_cabecera_salida.rows) > 2:
                        tabla_cabecera_salida._element.remove(tabla_cabecera_salida.rows[2]._element)

                    # Limpiar la tabla antes de agregar elementos de esta operación
                    while len(tabla_request.rows) > 2:
                        tabla_request._element.remove(tabla_request.rows[2]._element)
                        
                    # Limpiar la tabla antes de agregar elementos de esta operación
                    while len(tabla_response.rows) > 2:
                        tabla_response._element.remove(tabla_response.rows[2]._element)
                    
                    # Procesar los datos
                    for elem in elements['request']:
                        
                        obligatorio = "NO"
                        #if 'cabeceraEntrada.' not in elem['name']:
                        # Añadir una nueva fila al final de la tabla
                        #fila[0].text = operation + "Request" + "." + elem['name']
                        if 'cabeceraEntrada' in elem['name']:
                            fila_cabecera_entrada = tabla_cabecera_entrada.add_row().cells
                            fila_cabecera_entrada[0].text = elem['name']
                            #st.success(f"fila[0].text: {fila[0].text}")
                            fila_cabecera_entrada[1].text = elem['name']
                            campo = fila_cabecera_entrada[1].text.split('.')[-1]
                            fila_cabecera_entrada[1].text = campo
                            #st.success(f"fila[1].text: {fila[1].text}")
                            if elem['minOccurs'] == '1':
                                obligatorio = "SI"
                            fila_cabecera_entrada[2].text = obligatorio
                            fila_cabecera_entrada[3].text = elem['type']
                            tipo_campo = fila_cabecera_entrada[3].text.split(':')[-1]
                            if tipo_campo == 'string':
                                tipo_campo = 'String'
                            fila_cabecera_entrada[3].text = tipo_campo
                        else:
                            fila = tabla_request.add_row().cells
                            fila[0].text = elem['name']
                            #st.success(f"fila[0].text: {fila[0].text}")
                            fila[1].text = elem['name']
                            campo = fila[1].text.split('.')[-1]
                            fila[1].text = campo
                            #st.success(f"fila[1].text: {fila[1].text}")
                            if elem['minOccurs'] == '1':
                                obligatorio = "SI"
                            fila[2].text = obligatorio
                            fila[3].text = elem['type']
                            tipo_campo = fila[3].text.split(':')[-1]
                            if tipo_campo == 'string':
                                tipo_campo = 'String'
                            fila[3].text = tipo_campo
                        #st.success(f"fila[3].text: {fila[3].text}")
                    
                    if total_operaciones == 1:
                        progress_bar_general.progress(50)
                    
                    # Limpiar la tabla antes de agregar elementos de esta operación
                    while len(tabla_response.rows) > 2:
                        tabla_response._element.remove(tabla_response.rows[2]._element)
                    
                    # Procesar los datos
                    for elem in elements['response']:
                        
                        obligatorio = "NO"
                        #if 'cabeceraSalida.' not in elem['name']:
                        # Añadir una nueva fila al final de la tabla
                        # Rellenar la fila con los datos correspondientes
                        #fila[0].text = operation + "Response" + "." + elem['name']
                        if 'cabeceraSalida' in elem['name']:
                            fila_cabecera_salida = tabla_cabecera_salida.add_row().cells
                            fila_cabecera_salida[0].text = elem['name']
                            #st.success(f"fila[0].text: {fila[0].text}")
                            fila_cabecera_salida[1].text = elem['name']
                            campo = fila_cabecera_salida[1].text.split('.')[-1]
                            fila_cabecera_salida[1].text = campo
                            #st.success(f"fila[1].text: {fila[1].text}")
                            if elem['minOccurs'] == '1':
                                obligatorio = "SI"
                            fila_cabecera_salida[2].text = obligatorio
                            fila_cabecera_salida[3].text = elem['type']
                            tipo_campo = fila_cabecera_salida[3].text.split(':')[-1]
                            if tipo_campo == 'string':
                                tipo_campo = 'String'
                            fila_cabecera_salida[3].text = tipo_campo
                        else:
                            fila = tabla_response.add_row().cells
                            fila[0].text = elem['name']
                            #st.success(f"fila[0].text: {fila[0].text}")
                            fila[1].text = elem['name']
                            campo = fila[1].text.split('.')[-1]
                            fila[1].text = campo
                            #st.success(f"fila[1].text: {fila[1].text}")
                            if elem['minOccurs'] == '1':
                                obligatorio = "SI"
                            fila[2].text = obligatorio
                            fila[3].text = elem['type']
                            tipo_campo = fila[3].text.split(':')[-1]
                            if tipo_campo == 'string':
                                tipo_campo = 'String'
                            fila[3].text = tipo_campo
                    
                    if total_operaciones == 1:
                        progress_bar_general.progress(75)
                    
                    #st.success("___________________________________________")
                    
                    #st.success(f"✅ temp_dir  {temp_dir }")
                    #st.success(f"✅ ruta_temporal  {ruta_temporal }")

                    # Lista para almacenar las rutas de los documentos generados
                    documentos_generados = []

                    ruta_proyecto = ruta.strip("/")  # Asegurar que la ruta no tenga "/" al inicio
                    #st.success(f"✅ ruta_proyecto  {ruta_proyecto }")
                    nombre_documento = f"Especificación Servicio WSDL {operation}.docx"
                    
                    # Crear la ruta dentro de la carpeta temporal
                    carpeta_destino = os.path.join(ruta_temporal, ruta_proyecto)
                    os.makedirs(carpeta_destino, exist_ok=True)  # Crear la carpeta si no existe
                    
                    ruta_guardado = os.path.join(carpeta_destino, nombre_documento)
                    
                    doc_nuevo = replace_text_in_doc(doc, variables)
                    doc_nuevo.save(ruta_guardado)  # Guardar en la carpeta temporal
                    st.success(f"📄 Documento generado: ✅ {nombre_documento}")
                    
                    if total_operaciones == 1:
                        progress_bar_general.progress(100)
                    
                    
                    # 📌 Agregar el documento al ZIP
                    if os.path.exists(ruta_guardado):
                        zipf.write(ruta_guardado, os.path.join(ruta_proyecto, nombre_documento))
                        #st.success(f"📄 Documento agregado al ZIP: {ruta_guardado}")
                    else:
                        st.warning(f"⚠️ Documento no encontrado: {ruta_guardado}")
                    
                    generoArchivo = True
                        
        # 📥 Permitir la descarga del ZIP final
        with open(zip_path, "rb") as file:
            zip_bytes = file.read()
        
        progress_bar_general.progress(100)  # ¡Completado!
        st.success("Documentación generada con éxito!")

        # 🔹 Agregar un pequeño delay para asegurar que el ZIP esté listo
        time.sleep(2)  # Esperar 2 segundos antes de mostrar la descarga

        # 🔹 Descargar automáticamente el ZIP sin necesidad de clic
        st.download_button(
            label="📥 Descargar TODOS los documentos en ZIP",
            data=zip_bytes,
            file_name="Documentos_Completos.zip",
            mime="application/zip",
            key="download_all",
        )

def obtener_operaciones(project_path):

    operations =[]
    for root, dirs, files in os.walk(project_path):
        if os.path.basename(root) == "Proxies":
            ##st.success(f"✅ Proxies {elementos_xsd}")
            for file in files:
                if file.endswith('.ProxyService'):
                    osb_file_path = os.path.join(root, file)
                    #st.success(f"✅ osb_file_path {osb_file_path}")
                    project_name = extract_project_name_from_proxy(osb_file_path)
                    
                    if project_name is None:
                        continue 
                    pipeline_path = extract_pipeline_path_from_proxy(osb_file_path, project_path)
                    ##st.success(f"✅ pipeline_path {pipeline_path}")
                    with open(osb_file_path, 'r', encoding="utf-8") as f:
                        content = f.read()
                        if has_http_provider_id(content):
                            service_name = os.path.splitext(file)[0]
                            service_url = extract_service_url(content)
                            wsdl_relative_path = extract_wsdl_relative_path(content)
                            operacion_business = ""
                            if wsdl_relative_path:
                                wsdl_path = os.path.join(project_path, wsdl_relative_path + ".WSDL")
                                capa_proyecto = '/'+ wsdl_relative_path.split('/')[0]
                                
                                #st.success(f"capa_proyecto: {capa_proyecto}")
                                
                                #st.success(f"wsdl_path: {wsdl_path}")
                                operaciones_especificas = extract_wsdl_operations(wsdl_path)
                                #st.success(f"operations: {operations}")
                                
                                for operation in operaciones_especificas:
                                    operations.append(operation)
    return operations

def plantuml_to_hex(plantuml_code):
    hex_encoded = plantuml_code.encode("utf-8").hex()
    #print_with_line_number(f"hex_encoded: {hex_encoded}")
    return f"~h{hex_encoded}"  # Se agrega "~h" como indica la documentación

def encode_plantuml(text):
    """Codifica un diagrama de PlantUML en la versión comprimida para usar en URLs."""
    # Convertir a bytes en UTF-8
    data = text.encode("utf-8")
    
    # Comprimir con Deflate
    compressed_data = zlib.compress(data)[2:-4]  # Quitar cabecera y checksum
    
    # Convertir en Base64 modificada
    encoded = base64.b64encode(compressed_data).decode("utf-8")
    
    # Reemplazar caracteres según la tabla de PlantUML
    return encoded.translate(str.maketrans("ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789+/", PLANTUML_ENCODING))

def generate_plantuml_url(uml_text):
    """Genera una URL de PlantUML para visualizar el diagrama."""
    encoded_diagram = encode_plantuml(uml_text)
    return f"https://www.plantuml.com/plantuml/png/{encoded_diagram}"

def contiene_valor(valor_a_buscar, diccionario, profundidad=0, limite=50):
    if profundidad > limite:  # Limite de recursión para evitar desbordamiento
        return False
    
    for clave, valor in diccionario.items():
        if valor_a_buscar in clave:
            return True
        
        if isinstance(valor, str) and valor_a_buscar in valor:
            return True
        elif isinstance(valor, list):
            for item in valor:
                if isinstance(item, str) and valor_a_buscar in item:
                    return True
                elif isinstance(item, tuple):
                    if any(valor_a_buscar in str(subitem) for subitem in item):
                        return True
        elif isinstance(valor, dict):
            if contiene_valor(valor_a_buscar, valor, profundidad+1, limite):  # Aumentar profundidad
                return True
    
    return False

def descargar_diagrama(uml_url, ruta_destino):
    response = requests.get(uml_url)
    if response.status_code == 200:
        with open(ruta_destino, "wb") as file:
            file.write(response.content)
        print(f"Diagrama guardado en: {ruta_destino}")
        return ruta_destino
    else:
        print(f"Error al descargar diagrama: {response.status_code}")
        return None

def procesar_referencias_nivel2(referencia_padre,referencia_nueva,proxy, proxy_name, data, uml, profundidad=0):
    
    referencias_procesadas = set()
                
    print_with_line_number(f"♪EMPIEZA FLUJO procesar_referencias_nivel2 -> Referencia padre♪: {referencia_padre}")
    proyecto_padre = referencia_padre.split("/")[0]
    print_with_line_number(f"proyecto_padre procesar_referencias_nivel2: {proyecto_padre}")
    partes = referencia_nueva.split("/")
    if len(partes) >= 3:
        print_with_line_number(f"referencia_nueva procesar_referencias_nivel2: {referencia_nueva}")
        proyecto = partes[0]
        print_with_line_number(f"proyecto procesar_referencias_nivel2: {proyecto}")
        business = partes[1]
        print_with_line_number(f"business procesar_referencias_nivel2: {business}")
        proxy = partes[-1]
        print_with_line_number(f"proxy procesar_referencias_nivel2: {proxy}")
    
    referencia_key = f"REFERENCIA_{proxy}"
    
    if referencia_key in referencias_procesadas:
        return  # Ya fue procesado, evitamos duplicación
    
    referencias_procesadas.add(referencia_key)
    
    if referencia_key in data:
        print_with_line_number(f"{referencia_key} encontrado:")
        
        # 🔹 Obtener claves ordenadas (para saber cuál es la última)
        claves = list(data[referencia_key].keys())
        ultima_clave = claves[-1]  # Última clave en el diccionario
        print_with_line_number(f"🔽 Último elemento: {ultima_clave}")
        
        for key in claves:
            value = data[referencia_key][key]  # Valor de la clave
            print_with_line_number(f"value: {value}")
            division = value.split("/")
            project = division[0]
            print_with_line_number(f"procesar_referencias_nivel2 project: {project}")
            proyecto_business = division[1]
            print_with_line_number(f"procesar_referencias_nivel2 proyecto_business: {proyecto_business}")
            business_name = division[-1]

            print_with_line_number(f"procesar_referencias_nivel2 key - value {key}: {value}")

            if "ReglasNegocio" in value:
                regla_negocio = division[2]
                uml.append(f"{proxy_name} -> {regla_negocio}: Llamada a {business_name}")
                print_with_line_number(f"{proxy_name} -> {regla_negocio}: Llamada a {business_name}")
                uml.append(f"{regla_negocio} -> {proxy_name}: Retorna respuesta")
                print_with_line_number(f"{regla_negocio} -> {proxy_name}: Retorna respuesta")

            else:
                if "Proxies" in value:
                    nueva_referencia_key = f"REFERENCIA_{business_name}"
                    print_with_line_number(f"nueva_referencia_key: {nueva_referencia_key}")
                    
                    if not nueva_referencia_key in data:
                        uml.append(f"{proyecto} -> {project}: Llamada a {business_name}")
                        print_with_line_number(f"{proyecto} -> {project}: Llamada a {business_name}")
                        uml.append(f"{project} -> {proyecto}: Retorna respuesta")
                        print_with_line_number(f"{project} -> {proyecto}: Retorna respuesta")
                        if key == ultima_clave:
                            uml.append(f"{proyecto} -> {proxy_name}: Retorna respuesta")
                            print_with_line_number(f"{proyecto} -> {proxy_name}: Retorna respuesta")
                    else:
                        uml.append(f"{proyecto} -> {project}: Llamada a {business_name}")
                        print_with_line_number(f"{proyecto} -> {project}: Llamada a {business_name}")

                        print_with_line_number(f"value: {value}")
                        print_with_line_number(f"business_name: {business_name}")
                        print_with_line_number(f"project: {project}")
                        # #print_with_line_number(f"data: {data}")
                        
                        nueva_referencia_key = f"REFERENCIA_{business_name}"
                        print_with_line_number(f"nueva_referencia_key: {nueva_referencia_key}")
                    
                        if nueva_referencia_key in data:
                            print_with_line_number(f"{nueva_referencia_key} encontrado:")
                            claves_nuevas = list(data[nueva_referencia_key].keys())
                            ultima_clave_nueva = claves_nuevas[-1]  # Última clave en el diccionario
                            for key_nueva in claves_nuevas:
                                value_nuevo = data[nueva_referencia_key][key_nueva]  # Valor de la clave
                                print_with_line_number(f"procesar_referencias_nivel2 value_nuevo: {value_nuevo}")
                                partes_nuevas = value_nuevo.split("/")
                                project_nuevo = partes_nuevas[0]
                                print_with_line_number(f"procesar_referencias_nivel2 project_nuevo: {project_nuevo}")
                                proyecto_business_nuevo = partes_nuevas[1]
                                print_with_line_number(f"procesar_referencias_nivel2 proyecto_business_nuevo: {proyecto_business_nuevo}")
                                business_name_nuevo = partes_nuevas[-1]
                                print_with_line_number(f"procesar_referencias_nivel2 business_name_nuevo: {business_name_nuevo}")
                                
                        if key == ultima_clave:
                            uml.append(f"{project} -> {proxy_name}: Retorna respuesta")
                            print_with_line_number(f"{project} -> {proxy_name}: Retorna respuesta")
                        # 🔄 **Llamada recursiva**: buscamos si `business_name` también tiene una referencia

                else:
                    uml.append(f"{project} -> {proyecto_business}: Llamada a {business_name}")
                    print_with_line_number(f"{project} -> {proyecto_business}: Llamada a {business_name}")
                    uml.append(f"{proyecto_business} -> {project}: Retorna respuesta")
                    print_with_line_number(f"{proyecto_business} -> {project}: Retorna respuesta")

    else:
        
        if "BusinessServices" in referencia_nueva:
            uml.append(f"{proxy_name} -> {business}: Llamada a {proxy}")
            print_with_line_number(f"{proxy_name} -> {business}: Llamada a {proxy}")
            uml.append(f"{business} -> {proxy_name}: Retorna respuesta")
            print_with_line_number(f"{business} -> {proxy_name}: Retorna respuesta")
            if profundidad > 0:
                uml.append(f"{proxy_name} -> {proyecto_padre}: Retorna respuesta")
                print_with_line_number(f"{proxy_name} -> {proyecto_padre}: Retorna respuesta")
            else:
                uml.append(f"{proxy_name} -> {proyecto}: Retorna respuesta")
                print_with_line_number(f"{proxy_name} -> {proyecto}: Retorna respuesta")
            profundidad = 0



def generar_diagramas_operaciones(project_name, service_name, combined_services2, operacion_a_documentar=None):
    """
    Genera diagramas de secuencia para cada operación en combined_services2.
    """
    diagrama_path =""
    referencias_procesadas = set()
    proyecto_referencia_abc =""
    proxy_ebs = ""
    for operacion, detalles in combined_services2.items():
        
        if operacion_a_documentar == operacion:
        
            print_with_line_number(f"\n🔹 Operacion: {operacion}")
            
            uml = ["@startuml"]
            data = combined_services2[operacion]
        
            uml.append("skinparam maxMessageSize 270")
            # Lista para almacenar los participantes manteniendo el orden
            participantes = []
            
            def add_participant(alias, nombre):
                if (alias, nombre) not in participantes:
                    participantes.append((alias, nombre))
            
            
            
            
            def procesar_referencias(referencia_padre,referencia_nueva,proxy, proxy_name, data, uml, profundidad=0):
                
                print_with_line_number(f"♪EMPIEZA FLUJO -> Referencia padre♪: {referencia_padre}")
                proyecto_padre = referencia_padre.split("/")[0]
                print_with_line_number(f"proyecto_padre: {proyecto_padre}")
                partes = referencia_nueva.split("/")
                if len(partes) >= 3:
                    print_with_line_number(f"referencia_nueva: {referencia_nueva}")
                    proyecto = partes[0]
                    print_with_line_number(f"proyecto: {proyecto}")
                    business = partes[1]
                    print_with_line_number(f"business: {business}")
                    proxy = partes[-1]
                    print_with_line_number(f"proxy: {proxy}")
                
                referencia_key = f"REFERENCIA_{proxy}"
                
                if referencia_key in referencias_procesadas:
                    return  # Ya fue procesado, evitamos duplicación
                
                referencias_procesadas.add(referencia_key)
                
                if referencia_key in data:
                    print_with_line_number(f"{referencia_key} encontrado:")
                    
                    # 🔹 Obtener claves ordenadas (para saber cuál es la última)
                    claves = list(data[referencia_key].keys())
                    ultima_clave = claves[-1]  # Última clave en el diccionario
                    print_with_line_number(f"🔽 Último elemento: {ultima_clave}")
                    
                    for key in claves:
                        value = data[referencia_key][key]  # Valor de la clave
                        print_with_line_number(f"value: {value}")
                        division = value.split("/")
                        project = division[0]
                        print_with_line_number(f"project: {project}")
                        proyecto_business = division[1]
                        print_with_line_number(f"proyecto_business: {proyecto_business}")
                        business_name = division[-1]

                        print_with_line_number(f"key - value {key}: {value}")

                        if "ReglasNegocio" in value:
                            regla_negocio = division[2]
                            uml.append(f"{project} -> {regla_negocio}: Llamada a {business_name}")
                            print_with_line_number(f"{project} -> {regla_negocio}: Llamada a {business_name}")
                            uml.append(f"{regla_negocio} -> {project}: Retorna respuesta")
                            print_with_line_number(f"{regla_negocio} -> {project}: Retorna respuesta")

                        else:
                            if "Proxies" in value:
                                nueva_referencia_key = f"REFERENCIA_{business_name}"
                                print_with_line_number(f"nueva_referencia_key: {nueva_referencia_key}")
                                
                                if not nueva_referencia_key in data:
                                    uml.append(f"{proyecto} -> {project}: Llamada a {business_name}")
                                    print_with_line_number(f"{proyecto} -> {project}: Llamada a {business_name}")
                                    uml.append(f"{project} -> {proyecto}: Retorna respuesta")
                                    print_with_line_number(f"{project} -> {proyecto}: Retorna respuesta")
                                    if key == ultima_clave:
                                        uml.append(f"{proyecto} -> {proxy_name}: Retorna respuesta")
                                        print_with_line_number(f"{proyecto} -> {proxy_name}: Retorna respuesta")
                                else:
                                    uml.append(f"{proyecto} -> {project}: Llamada a {business_name}")
                                    print_with_line_number(f"{proyecto} -> {project}: Llamada a {business_name}")

                                    print_with_line_number(f"value: {value}")
                                    print_with_line_number(f"business_name: {business_name}")
                                    print_with_line_number(f"project: {project}")
                                    # #print_with_line_number(f"data: {data}")
                                    
                                    referencia_nueva_key = f"REFERENCIA_{business_name}"
                                    print_with_line_number(f"nueva_referencia_key: {referencia_nueva_key}")
                                
                                    if referencia_nueva_key in data:
                                        print_with_line_number(f"{referencia_nueva_key} encontrado:")
                                        claves_nuevas = list(data[referencia_nueva_key].keys())
                                        ultima_clave_nueva = claves_nuevas[-1]  # Última clave en el diccionario
                                        for key_nueva in claves_nuevas:
                                            value_nuevo = data[referencia_nueva_key][key_nueva]  # Valor de la clave
                                            print_with_line_number(f"value_nuevo: {value_nuevo}")
                                            partes_nuevas = value_nuevo.split("/")
                                            project_nuevo = partes_nuevas[0]
                                            print_with_line_number(f"project_nuevo: {project_nuevo}")
                                            proyecto_business_nuevo = partes_nuevas[1]
                                            print_with_line_number(f"proyecto_business_nuevo: {proyecto_business_nuevo}")
                                            business_name_nuevo = partes_nuevas[-1]
                                            print_with_line_number(f"business_name_nuevo: {business_name_nuevo}")
                                            
                                            nueva_referencia_business_key = f"REFERENCIA_{business_name_nuevo}"
                                            print_with_line_number(f"nueva_referencia_key: {nueva_referencia_business_key}")
                                            project_business = value_nuevo.split("/")[1]
                                            
                                            if nueva_referencia_business_key in data:
                                                print_with_line_number(f"{nueva_referencia_business_key} encontrado:")
                                                claves_nuevas = list(data[nueva_referencia_business_key].keys())
                                                ultima_clave_nueva = claves_nuevas[-1] 
                                                
                                                uml.append(f"{project} -> {project_nuevo}: Llamada a {business_name_nuevo}")
                                                print_with_line_number(f"{project} -> {project_nuevo}: Llamada a {business_name_nuevo}")
                                                
                                                procesar_referencias(referencia_padre,value_nuevo,business_name_nuevo, project_nuevo, data, uml, profundidad + 1)
                                            else:
                                                if "Proxies" in value_nuevo:
                                                    uml.append(f"{project} -> {project_nuevo}: Llamada a {business_name_nuevo}")
                                                    print_with_line_number(f"{project} -> {project_nuevo}: Llamada a {business_name_nuevo}")
                                                    uml.append(f"{project_nuevo} -> {project}: Retorna respuesta")
                                                    print_with_line_number(f"{project_nuevo} -> {project}: Retorna respuesta")
                                                    
                                                    if key_nueva == ultima_clave_nueva:
                                                        uml.append(f"{project} -> {proyecto}: Retorna respuesta")
                                                        print_with_line_number(f"{project} -> {proyecto}: Retorna respuesta")
                                                else:
                                                    uml.append(f"{project} -> {project_business}: Llamada a {business_name_nuevo}")
                                                    print_with_line_number(f"{project} -> {project_business}: Llamada a {business_name_nuevo}")
                                                    uml.append(f"{project_business} -> {project}: Retorna respuesta")
                                                    print_with_line_number(f"{project_business} -> {project}: Retorna respuesta")
                                                    uml.append(f"{project} -> {proyecto}: Retorna respuesta")
                                                    print_with_line_number(f"{project} -> {proyecto}: Retorna respuesta")
                                           
                                    # if key == ultima_clave:
                                        # uml.append(f"{project} -> {proxy_name[0]}: Retorna respuesta")
                                        # print_with_line_number(f"{project} -> {proxy_name}: Retorna respuesta")
                                    # # 🔄 **Llamada recursiva**: buscamos si `business_name` también tiene una referencia
                                    # #procesar_referencias(referencia_padre,value,business_name, project, data, uml, profundidad + 1)

                            else:
                                uml.append(f"{project} -> {proyecto_business}: Llamada a {business_name}")
                                print_with_line_number(f"{project} -> {proyecto_business}: Llamada a {business_name}")
                                uml.append(f"{proyecto_business} -> {project}: Retorna respuesta")
                                print_with_line_number(f"{proyecto_business} -> {project}: Retorna respuesta")

                else:
                    
                    if "BusinessServices" in referencia_nueva:
                        uml.append(f"{proxy_name} -> {business}: Llamada a {proxy}")
                        print_with_line_number(f"{proxy_name} -> {business}: Llamada a {proxy}")
                        uml.append(f"{business} -> {proxy_name}: Retorna respuesta")
                        print_with_line_number(f"{business} -> {proxy_name}: Retorna respuesta")
                        if profundidad > 0:
                            uml.append(f"{proxy_name} -> {proyecto_padre}: Retorna respuesta")
                            print_with_line_number(f"{proxy_name} -> {proyecto_padre}: Retorna respuesta")
                        else:
                            uml.append(f"{proxy_name} -> {proyecto}: Retorna respuesta")
                            print_with_line_number(f"{proxy_name} -> {proyecto}: Retorna respuesta")
                        profundidad = 0
            
            
            
            # Agregar siempre el usuario y el EXP con el nombre dinámico
            add_participant("Usuario", "Usuario")
            add_participant("EXP", project_name)
            
            # Verificar los datos en la estructura y agregar solo si no existen
            if "Proxy" in data:
                for proxy in data["Proxy"]:
                    proyecto_ebs = proxy.split("/")[0]
                    add_participant(proyecto_ebs, proyecto_ebs)
            
            print_with_line_number(f"proyecto_ebs: {proyecto_ebs}")
            print_with_line_number(f"participantes: {participantes}")
        
            if contiene_valor("ReglasNegocio",data):
                print_with_line_number(f"Existe ReglasNegocio")
                add_participant("ReglasNegocio", "ReglasNegocio")
            if contiene_valor("BPEL",data):
                print_with_line_number(f"Existe BPEL")
                add_participant("BPEL", "BPEL")
                print_with_line_number(f"Existe BPEL")
            
            print_with_line_number(f"participantes: {participantes}")
        
            if "Referencia" in data:
                for referencia in data["Referencia"]:
                    partes = referencia.split("/")
                    if len(partes) >= 3:
                        proyecto = partes[0]
                        business = partes[1]
                        proxy = partes[-1]
                        print_with_line_number(f"Proyecto: {proyecto}, Business: {business}, Proxy: {proxy}")
                        add_participant(proyecto, proyecto)
                        if "BusinessServices" in business:
                            add_participant(business, business)
            
            print_with_line_number(f"participantes: {participantes}")
            
            if any(key.startswith("REFERENCIA_") for key in data):
                for key in data:
                    print_with_line_number(f"key: {key}")
                    if key.startswith("REFERENCIA_"):
                        for sub_ref in data[key]:
                            print_with_line_number(f"sub_ref: {sub_ref}")
                            clave = data[key][sub_ref]
                            print_with_line_number(f"clave: {clave}")
                            if "BusinessServices" in clave:
                                business = clave.split("/")[1]
                                add_participant(business, business)
                            ref_name = data[key][sub_ref].split("/")[0]
                            print_with_line_number(f"ref_name: {ref_name}")
                            add_participant(ref_name, ref_name)
            
            print_with_line_number(f"participantes: {participantes}")
            
            # Agregar los participantes al diagrama
            for alias, nombre in participantes:
                if alias == "BusinessServices":
                    uml.append(f"database {nombre} as {alias}")
                else:
                    uml.append(f"participant {nombre} as {alias}")
            
            business_services = 'database BusinessServices as BusinessServices'
            # Si 'BusinessServices' está en la lista, lo mueve al final
            if business_services in uml_elements:
                uml_elements.remove(business_services)  # Elimina si existe
                uml_elements.append(business_services)  # Lo agrega al final
            print_with_line_number(f"uml: {uml}")
            
            uml.append(f"Usuario -[#red]> EXP: Llamada a {operacion} en {service_name}")
            if "Proxy" in data:
                for proxy in data["Proxy"]:
                    proxy_ebs = proxy
                    proyecto_ebs = proxy.split("/")[0]
                    uml.append(f"EXP -> {proyecto_ebs}: Llamada a {proxy.split('/')[-1]}")
            print_with_line_number(f"uml: {uml}")
            
            proyecto_referencia_abc =""
            if "Referencia" in data:
                for referencia in data["Referencia"]:
                    partes = referencia.split("/")
                    if len(partes) >= 3:
                        print_with_line_number(f"referencia: {referencia}")
                        proyecto_referencia_abc = partes[0]
                        print_with_line_number(f"proyecto_referencia_abc: {proyecto_referencia_abc}")
                        business = partes[1]
                        print_with_line_number(f"business: {business}")
                        proxy = partes[-1]
                        print_with_line_number(f"proxy : {proxy}")
                        
                        if proyecto_ebs != proyecto_referencia_abc:
                            uml.append(f"{proyecto_ebs} -> {proyecto_referencia_abc}: Llamada a {proxy}")
                            print_with_line_number(f"{proyecto_ebs} -> {proyecto_referencia_abc}: Llamada a {proxy}")
                        if "BusinessServices" in business:
                            uml.append(f"{proyecto_referencia_abc} -> {business}: Llamada a {proxy}")
                            print_with_line_number(f"{proyecto_referencia_abc} -> {business}: Llamada a {proxy}")
                            uml.append(f"{business} -> {proyecto_referencia_abc}: Retorna respuesta")
                            print_with_line_number(f"{business} -> {proyecto_referencia_abc}: Retorna respuesta")
                        else:
                            procesar_referencias(proxy_ebs,referencia,proxy, proxy_ebs, data, uml)
                            uml.append(f"{proyecto_referencia_abc} -> {proyecto_ebs}: Retorna respuesta")
                            print_with_line_number(f"{proyecto_referencia_abc} -> {proyecto_ebs}: Retorna respuesta")
                            
                            # referencia_key = f"REFERENCIA_{proxy}"
                            # if referencia_key in data:
                                # print_with_line_number(f"  - {referencia_key} encontrado:")
                                # for key, value in data[referencia_key].items():
                                    # project = value.split("/")[0]
                                    # proyecto_business = value.split("/")[1]
                                    # business_name = value.split("/")[-1]
                                    # print_with_line_number(f"referencia_key: {referencia_key}")
                                    # print_with_line_number(f"proyecto_business: {proyecto_business}")
                                    # print_with_line_number(f"    - {key}: {value}")
                                    # if "ReglasNegocio" in value:
                                        # regla_negocio = value.split("/")[2]
                                        # uml.append(f"{proxy_name} -> {regla_negocio}: Llamada a {business_name}")
                                        # print_with_line_number(f"{proxy_name} -> {regla_negocio}: Llamada a {business_name}")
                                        # uml.append(f"{regla_negocio} -> {proxy_name}: Retorna respuesta")
                                        # print_with_line_number(f"{regla_negocio} -> {proxy_name}: Retorna respuesta")
                                    # else:
                                        # if "Proxies" in proyecto_business:
                                            # uml.append(f"{proxy_name} -> {proyecto}: Llamada a {proxy}")
                                            # print_with_line_number(f"{proxy_name} -> {proyecto}: Llamada a {proxy}")
                                            # uml.append(f"{proyecto} -> {project}: Llamada a {business_name}")
                                            # print_with_line_number(f"{proyecto} -> {project}: Llamada a {business_name}")
                                            # uml.append(f"{project} -> {proyecto}: Retorna respuesta")
                                            # print_with_line_number(f"{project} -> {proyecto}: Retorna respuesta")
                                            # uml.append(f"{proyecto} -> {proxy_name}: Retorna respuesta")
                                            # print_with_line_number(f"{proyecto} -> {proxy_name}: Retorna respuesta")

                                        # else:
                                            # uml.append(f"{proxy_name} -> {proyecto}: Llamada a {proxy}")
                                            # print_with_line_number(f"{proxy_name} -> {proyecto}: Llamada a {proxy}")
                                            # uml.append(f"{proyecto} -> {proyecto_business}: Llamada a {business_name}")
                                            # print_with_line_number(f"{proyecto} -> {proyecto_business}: Llamada a {business_name}")
                                            # uml.append(f"{proyecto_business} -> {proyecto}: Retorna respuesta")
                                            # print_with_line_number(f"{proyecto_business} -> {proyecto}: Retorna respuesta")
                                            # uml.append(f"{proyecto} -> {proxy_name}: Retorna respuesta")
                                            # print_with_line_number(f"{proyecto} -> {proxy_name}: Retorna respuesta")
                            # else:
                                # uml.append(f"{proxy_name} -> {proyecto}: Llamada a {proxy}")
                                # print_with_line_number(f"{proxy_name} -> {proyecto}: Llamada a {proxy}")
                                # uml.append(f"{proyecto} -> {proxy_name}: Retorna respuesta")
                                # print_with_line_number(f"{proyecto} -> {proxy_name}: Retorna respuesta")
                                # #uml.append(f"{proxy_name} -> EXP: Retorna respuesta")
                                # print_with_line_number(f"{proxy_name} -> EXP: Retorna respuesta")
                        
                uml.append(f"{proyecto_ebs} -> EXP: Retorna respuesta")
                print_with_line_number(f"{proyecto_ebs} -> EXP: Retorna respuesta")
            print_with_line_number(f"uml: {uml}")
            
            uml.append("EXP -> Usuario : Respuesta final")
            uml.append("@enduml")
        
            print_with_line_number("\n".join(uml))
        
            encoded_code = plantuml_to_hex("\n".join(uml))
            img_url = f"{PLANTUML_SERVER}{encoded_code}"    
            
            #st.image(img_url, caption=f"Diagrama de {operacion}", use_container_width=True)
            #st.markdown(f"[Descargar {operacion}]({img_url})", unsafe_allow_html=True)
            
            # Generar URL
            uml_url = generate_plantuml_url("\n".join(uml))
            #print_with_line_number(f"URL del diagrama: {uml_url}")
            
            # URL final
            plantuml_url_png = {uml_url}
            #st.image(plantuml_url_png)
            #print("🔹 URL de la imagen PNG:", plantuml_url_png)
            
            plantuml_url_png = uml_url
            # Descargar la imagen del servidor de PlantUML
            output_dir = "diagramas"
            os.makedirs(output_dir, exist_ok=True)
            diagrama_path = os.path.join(output_dir, f"{project_name}_{operacion}.png")
            #print_with_line_number(f"diagrama_path: {diagrama_path}")
            
            try:
                response = requests.get(plantuml_url_png)
                if response.status_code == 200:
                    with open(diagrama_path, "wb") as file:
                        file.write(response.content)
                        #print_with_line_number(f"Se guardo imagen en: {diagrama_path}")
                else:
                    print_with_line_number(f"Error al generar el diagrama: {response.status_code}")
            except Exception as e:
                print_with_line_number(f"Error en la solicitud de la imagen: {e}")
    
    return diagrama_path
    

def main():
    st.markdown(
    "<h1 style='text-align: center;'>📄 Generador de Documentación OSB</h1>",
    unsafe_allow_html=True)
    
    # Ruta donde se extraerán los archivos
    carpeta_destino = "extraccion_jar"
    operacion_a_documentar = ""
    
    # 📌 Agregar elementos al menú lateral
    with st.sidebar:
        jar_file = st.file_uploader("Sube el archivo .jar con dependencias", type=["jar"])
        plantilla_file = st.file_uploader("Sube la plantilla de Word", type=["docx"])
        if jar_file:
            jar_path = "temp.jar"

            # 🔥 Borrar contenido previo de la carpeta `extraccion_jar` solo si existe
            if os.path.exists(carpeta_destino):
                try:
                    shutil.rmtree(carpeta_destino)  # Elimina la carpeta y su contenido
                except Exception as e:
                    st.success(f"⚠️ No se pudo limpiar la carpeta temporal: {e}")

            # 📌 Crear nuevamente la carpeta vacía
            os.makedirs(carpeta_destino, exist_ok=True)

            # Guardar el nuevo archivo .jar
            with open(jar_path, "wb") as f:
                f.write(jar_file.getbuffer())

            # 📂 Extraer los archivos del nuevo .jar
            try:
                with zipfile.ZipFile(jar_path, "r") as jar:
                    jar.extractall(carpeta_destino)
                    archivos_extraidos = jar.namelist()

                #st.success(f"✅ Archivos extraídos en: {carpeta_destino}")
            except zipfile.BadZipFile:
                st.error("❌ Error: El archivo no es un JAR válido o está dañado.")
            
            operaciones = obtener_operaciones(carpeta_destino)
            # Agregar una opción vacía al inicio de la lista
            operaciones.insert(0, "TODAS")
            if operaciones:  # Solo mostrar si hay operaciones disponibles
                operacion_a_documentar = st.selectbox("Selecciona una operación", operaciones)
                if operacion_a_documentar == "TODAS":
                    operacion_a_documentar = None
            else:
                st.warning("⚠️ No se encontraron operaciones disponibles.")
                operacion_a_documentar = None  # Para evitar errores si está vacío           
        nombre_autor = st.text_input("Nombre del autor", value="Kevin Torres")  # Valor por defecto
        generar_doc = st.button("Generar Documentación")
         
    with st.container():
        if generar_doc:
            if jar_file and plantilla_file and nombre_autor:
                #st.success(f"✅ operacion_a_documentar: {operacion_a_documentar}")
                with st.spinner("Generando documentación..."):
                    generar_documentacion(carpeta_destino, plantilla_file,operacion_a_documentar,nombre_autor)
            else:
                st.error("Por favor, sube todos los archivos, escribe el autor y sube la plantilla.")
                

if __name__ == "__main__":
    main()
